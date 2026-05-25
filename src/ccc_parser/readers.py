from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
from dataclasses import dataclass, field
from email import policy
from email.parser import BytesParser
from html import unescape
from pathlib import Path
from xml.etree import ElementTree as ET

from .models import ImageRecord, ParserWarning, SourceFile
from .normalization import clean_value
from .triage import file_sha256

try:  # pragma: no cover - optional import availability varies by machine
    import fitz
except Exception:  # pragma: no cover
    fitz = None

try:  # pragma: no cover
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:  # pragma: no cover
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:  # pragma: no cover
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:  # pragma: no cover
    import extract_msg
except Exception:  # pragma: no cover
    extract_msg = None

try:  # pragma: no cover
    from PIL import Image as PILImage
except Exception:  # pragma: no cover
    PILImage = None

try:  # pragma: no cover
    import pytesseract
except Exception:  # pragma: no cover
    pytesseract = None

try:  # pragma: no cover
    import olefile
except Exception:  # pragma: no cover
    olefile = None


OCR_PAGE_LIMIT = 2
SUBPROCESS_TIMEOUT_SECONDS = 30
ATTACHMENT_CACHE = Path(tempfile.gettempdir()) / "ccc_parser_attachments"
_TESSERACT_CONFIGURED: bool | None = None


@dataclass(frozen=True)
class DocumentModel:
    source_file: SourceFile
    text: str
    document_type: str
    reader_method: str
    reader_notes: list[ParserWarning] = field(default_factory=list)
    images: list[ImageRecord] = field(default_factory=list)
    attachments: list[Path] = field(default_factory=list)


def _strip_html_tags(value: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?</\1>", " ", value or "")
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p\s*>", "\n", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _safe_attachment_name(name: str, fallback: str) -> str:
    candidate = Path(name or fallback).name.strip() or fallback
    candidate = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", candidate)
    return candidate[:180] or fallback


def _attachment_output_dir(message_path: Path) -> Path:
    digest = file_sha256(message_path)[:16]
    output_dir = ATTACHMENT_CACHE / f"{digest}-{os.getpid()}-{uuid.uuid4().hex[:12]}"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _unique_attachment_path(output_dir: Path, filename: str) -> Path:
    stem = Path(filename).stem or "attachment"
    suffix = Path(filename).suffix
    candidate = output_dir / filename
    index = 2
    while candidate.exists():
        candidate = output_dir / f"{stem}-{index}{suffix}"
        index += 1
    return candidate


def _resource_path(name: str) -> Path:
    if hasattr(sys, "_MEIPASS"):
        return Path(getattr(sys, "_MEIPASS")) / name
    return Path(__file__).resolve().parents[2] / name


def _configure_tesseract() -> bool:
    global _TESSERACT_CONFIGURED
    if _TESSERACT_CONFIGURED is not None:
        return _TESSERACT_CONFIGURED
    if pytesseract is None:
        _TESSERACT_CONFIGURED = False
        return False

    explicit = os.environ.get("CCC_TESSERACT_CMD") or os.environ.get("TESSERACT_CMD")
    command_candidates = [Path(explicit)] if explicit else []
    command_candidates.extend(
        [
            _resource_path("tesseract") / "tesseract.exe",
            _resource_path("tesseract") / "tesseract",
            Path(__file__).resolve().parents[3] / "cedocumentmapper" / "tesseract" / "tesseract.exe",
            Path(__file__).resolve().parents[3] / "cedocumentmapper" / "tesseract" / "tesseract",
        ]
    )
    path_tesseract = shutil.which("tesseract")
    if path_tesseract:
        command_candidates.append(Path(path_tesseract))

    binary = next((candidate for candidate in command_candidates if candidate and candidate.exists()), None)
    if binary is None:
        _TESSERACT_CONFIGURED = False
        return False

    try:
        pytesseract.pytesseract.tesseract_cmd = str(binary)
    except Exception:
        _TESSERACT_CONFIGURED = False
        return False
    tessdata = binary.parent / "tessdata"
    if tessdata.exists():
        os.environ.setdefault("TESSDATA_PREFIX", str(tessdata))
    _TESSERACT_CONFIGURED = True
    return True


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except Exception:
            continue
    return path.read_text(errors="ignore")


def _decode_pypdf_unicode_sequences(value: str) -> str:
    return re.sub(r"/uni([0-9A-Fa-f]{4})", lambda match: chr(int(match.group(1), 16)), value or "")


def _extract_pdf_pymupdf(path: Path) -> tuple[str, list[ParserWarning], str]:
    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed.")
    warnings: list[ParserWarning] = []
    text_parts: list[str] = []
    image_counts: list[int] = []
    doc = fitz.open(path)
    try:
        for page_index, page in enumerate(doc, start=1):
            try:
                image_counts.append(len(page.get_images() or []))
            except Exception:
                image_counts.append(0)
            blocks = page.get_text("blocks", sort=True) or []
            page_blocks: list[str] = []
            for block in blocks:
                if len(block) >= 7 and block[6] != 0:
                    continue
                raw = str(block[4] if len(block) > 4 else "").replace("\r", "\n")
                lines = [line.rstrip() for line in raw.splitlines() if line.strip()]
                if lines:
                    page_blocks.append("\n".join(lines))
            if page_blocks:
                text_parts.append("\n\n".join(page_blocks))
                continue
            fallback = (page.get_text("text", sort=True) or "").replace("\r", "\n").strip()
            if fallback:
                text_parts.append(fallback)
            else:
                warnings.append(
                    ParserWarning(
                        code="pdf_page_no_text",
                        message=f"Page {page_index} had no selectable text.",
                    )
                )
        combined = "\n\n".join(part for part in text_parts if part.strip()).strip()
        should_ocr = (
            not combined
            and 0 < len(image_counts) <= OCR_PAGE_LIMIT
            and all(count == 1 for count in image_counts)
            and _configure_tesseract()
            and PILImage is not None
        )
        if should_ocr:
            ocr_pages: list[str] = []
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72))
                image = PILImage.open(io.BytesIO(pix.tobytes("png")))
                page_text = pytesseract.image_to_string(image, lang="eng") or ""
                if page_text.strip():
                    ocr_pages.append(page_text.replace("\r", "\n").strip())
            if ocr_pages:
                warnings = [
                    ParserWarning(
                        code="ocr_fallback_used",
                        message="Read PDF using OCR fallback because no native text layer was available.",
                    )
                ]
                return "\n\n".join(ocr_pages).strip(), warnings, "pymupdf_ocr_fallback"
    finally:
        doc.close()
    if not combined:
        raise RuntimeError("PyMuPDF returned no text.")
    return combined, warnings, "pymupdf_blocks"


def _extract_pdf_pdfplumber(path: Path) -> tuple[str, list[ParserWarning], str]:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is not installed.")
    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            table_lines: list[str] = []
            for table in tables:
                for row in table:
                    values = [clean_value(cell or "") for cell in row if clean_value(cell or "")]
                    if values:
                        table_lines.append(" | ".join(values))
            parts = [part for part in (text.strip(), "\n".join(table_lines).strip()) if part]
            if parts:
                text_parts.append("\n".join(parts))
    if not text_parts:
        raise RuntimeError("pdfplumber returned no text.")
    return "\n\n".join(text_parts), [], "pdfplumber_text_tables"


def _extract_pdf_pypdf(path: Path) -> tuple[str, list[ParserWarning], str]:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed.")
    text_parts: list[str] = []
    warnings: list[ParserWarning] = []
    reader = PdfReader(path)
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = _decode_pypdf_unicode_sequences(page.extract_text() or "")
        if page_text.strip():
            text_parts.append(page_text)
        else:
            warnings.append(ParserWarning(code="pdf_page_no_text", message=f"Page {page_index} had no selectable text."))
    if not text_parts:
        raise RuntimeError("pypdf returned no text.")
    return "\n\n".join(text_parts), warnings, "pypdf_text"


def _pdf_page_image_records(path: Path, source_file: SourceFile) -> list[ImageRecord]:
    if fitz is None:
        return []
    records: list[ImageRecord] = []
    try:
        doc = fitz.open(path)
    except Exception:
        return records
    try:
        for page_index, page in enumerate(doc, start=1):
            image_count = len(page.get_images() or [])
            if image_count < 1:
                continue
            rect = page.rect
            records.append(
                ImageRecord(
                    image_id=f"image-{source_file.sha256[:12]}-p{page_index:03d}",
                    source_file_id=source_file.file_id,
                    path=source_file.path,
                    sha256=source_file.sha256,
                    width=int(rect.width),
                    height=int(rect.height),
                    role_guess="pdf_page_image",
                    order_hint=page_index,
                )
            )
    finally:
        doc.close()
    return records


def _extract_docx_header_footer_lines(path: Path) -> tuple[list[str], list[str]]:
    def dedupe(lines: list[str]) -> list[str]:
        output: list[str] = []
        seen: set[str] = set()
        for line in lines:
            key = line.lower()
            if key not in seen:
                output.append(line)
                seen.add(key)
        return output

    def extract_part_lines(zf: zipfile.ZipFile, prefix: str) -> list[str]:
        collected: list[str] = []
        names = [
            name
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml") and Path(name).name.lower().startswith(prefix)
        ]
        for name in names:
            try:
                root = ET.fromstring(zf.read(name))
            except Exception:
                continue
            parts: list[str] = []
            for node in root.iter():
                tag_name = node.tag.rsplit("}", 1)[-1].lower()
                if tag_name == "t" and node.text:
                    parts.append(node.text)
                elif tag_name == "tab":
                    parts.append("\t")
                elif tag_name in {"br", "cr", "p"}:
                    value = clean_value("".join(parts))
                    if value:
                        collected.append(value)
                    parts = []
            value = clean_value("".join(parts))
            if value:
                collected.append(value)
        return dedupe(collected)

    try:
        with zipfile.ZipFile(path) as zf:
            return extract_part_lines(zf, "header"), extract_part_lines(zf, "footer")
    except Exception:
        return [], []


def _extract_docx_textbox_lines(path: Path) -> list[str]:
    output: list[str] = []
    try:
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                try:
                    root = ET.fromstring(zf.read(name))
                except Exception:
                    continue
                for node in root.iter():
                    if node.tag.rsplit("}", 1)[-1].lower() not in {"txbxcontent", "textbox"}:
                        continue
                    parts: list[str] = []
                    for child in node.iter():
                        tag_name = child.tag.rsplit("}", 1)[-1].lower()
                        if tag_name == "t" and child.text:
                            parts.append(child.text)
                        elif tag_name in {"p", "br", "cr"} and parts:
                            output.append(clean_value(" ".join(parts)))
                            parts = []
                    if parts:
                        output.append(clean_value(" ".join(parts)))
    except Exception:
        return []
    seen: set[str] = set()
    cleaned: list[str] = []
    for line in output:
        if line and line.lower() not in seen:
            cleaned.append(line)
            seen.add(line.lower())
    return cleaned


def _extract_docx_text(path: Path) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")
    doc = Document(path)
    parts: list[str] = []
    seen: set[str] = set()

    def append(value: str) -> None:
        value = (value or "").rstrip()
        if not value:
            parts.append("")
            return
        key = value.strip().lower()
        if key not in seen:
            parts.append(value)
            seen.add(key)

    headers, footers = _extract_docx_header_footer_lines(path)
    for line in headers:
        append(line)
    if headers:
        append("")
    for paragraph in doc.paragraphs:
        append(paragraph.text or "")
    for table in doc.tables:
        append("")
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                append(" | ".join(values))
    textboxes = _extract_docx_textbox_lines(path)
    if textboxes:
        append("")
        append("[Text Boxes]")
        for line in textboxes:
            append(line)
    if footers:
        append("")
        for line in footers:
            append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(parts)).strip("\n")


def _unique_story_lines(raw_text: str) -> list[str]:
    raw_text = (raw_text or "").replace("\r\x07", "\n").replace("\x07", "").replace("\r", "\n")
    output: list[str] = []
    seen: set[str] = set()
    for raw_line in raw_text.splitlines():
        line = clean_value(raw_line)
        if line and line.lower() not in seen:
            output.append(line)
            seen.add(line.lower())
    return output


def _extract_doc_text_via_word(path: Path) -> str:
    try:
        import pythoncom  # type: ignore
        from win32com.client import DispatchEx  # type: ignore
    except Exception as exc:
        raise RuntimeError("Microsoft Word automation is not available.") from exc

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(
            str(path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )
        headers: list[str] = []
        footers: list[str] = []
        for section_index in range(1, doc.Sections.Count + 1):
            section = doc.Sections(section_index)
            for hf_type in (1, 2, 3):
                try:
                    headers.extend(_unique_story_lines(section.Headers(hf_type).Range.Text or ""))
                except Exception:
                    pass
                try:
                    footers.extend(_unique_story_lines(section.Footers(hf_type).Range.Text or ""))
                except Exception:
                    pass
        content = (doc.Content.Text or "").replace("\r\x07", "\n").replace("\x07", "").replace("\r", "\n")
        parts = ["\n".join(headers), content.strip(), "\n".join(footers)]
        return "\n\n".join(part for part in parts if part.strip()).strip()
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _extract_rtf_text(path: Path) -> str:
    data = path.read_bytes()
    if not data.lstrip().startswith(b"{\\rtf"):
        raise RuntimeError("File is not RTF content.")
    text = data.decode("cp1252", errors="ignore")
    text = re.sub(
        r"\\'([0-9A-Fa-f]{2})",
        lambda match: bytes([int(match.group(1), 16)]).decode("cp1252", errors="ignore"),
        text,
    )

    def unicode_repl(match: re.Match[str]) -> str:
        value = int(match.group(1))
        if value < 0:
            value += 65536
        try:
            return chr(value)
        except ValueError:
            return ""

    text = re.sub(r"\\u(-?\d+)\??", unicode_repl, text)
    text = re.sub(r"\\(?:par|line|tab)\b[ ]?", "\n", text)
    text = re.sub(r"{\\fonttbl.*?}", " ", text, flags=re.DOTALL)
    text = re.sub(r"{\\colortbl.*?}", " ", text, flags=re.DOTALL)
    text = re.sub(r"{\\stylesheet.*?}", " ", text, flags=re.DOTALL)
    text = re.sub(r"{\\\*.*?}", " ", text, flags=re.DOTALL)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?|\\[^a-zA-Z]", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"[ \t]{2,}", " ", text)
    lines = [clean_value(line) for line in text.splitlines()]
    lines = [line for line in lines if line and not _looks_like_binary_noise(line)]
    if not lines:
        raise RuntimeError("No readable RTF text was found.")
    return "\n".join(lines)


def _extract_doc_text_via_soffice(path: Path) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice is not installed.")
    with tempfile.TemporaryDirectory() as tmpdir:
        outdir = Path(tmpdir)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", str(outdir), str(path.resolve())],
                check=True,
                capture_output=True,
                timeout=SUBPROCESS_TIMEOUT_SECONDS,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("LibreOffice DOC conversion timed out.") from exc
        candidates = list(outdir.glob("*.docx"))
        if not candidates:
            raise RuntimeError("LibreOffice did not produce a DOCX file.")
        return _extract_docx_text(candidates[0])


def _extract_doc_text_via_antiword(path: Path) -> str:
    antiword = shutil.which("antiword")
    if not antiword:
        raise RuntimeError("antiword is not installed.")
    try:
        result = subprocess.run(
            [antiword, str(path.resolve())],
            check=True,
            capture_output=True,
            timeout=SUBPROCESS_TIMEOUT_SECONDS,
        )
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("antiword DOC extraction timed out.") from exc
    output = result.stdout.decode("utf-8", errors="ignore")
    output = output.replace("\r\n", "\n").replace("\r", "\n")
    output = re.sub(r"\n{3,}", "\n\n", output)
    output = re.sub(r"[ \t]{2,}", " ", output)
    return output.strip()


def _extract_doc_text_via_olefile(path: Path) -> str:
    if olefile is None:
        raise RuntimeError("olefile is not installed.")
    ole = olefile.OleFileIO(path)
    try:
        stream_names = {"/".join(parts): parts for parts in ole.listdir()}
        candidate_parts: list[str] = []
        for stream_name in ("WordDocument", "1Table", "0Table"):
            if stream_name not in stream_names:
                continue
            data = ole.openstream(stream_names[stream_name]).read()
            for encoding in ("cp1252", "latin-1", "utf-16le"):
                decoded = data.decode(encoding, errors="ignore")
                candidate_parts.extend(_printable_doc_lines(decoded))
        lines: list[str] = []
        seen: set[str] = set()
        for line in candidate_parts:
            key = line.lower()
            if key in seen:
                continue
            if _looks_like_binary_noise(line):
                continue
            lines.append(line)
            seen.add(key)
        if not lines:
            raise RuntimeError("No printable WordDocument text was found.")
        return "\n".join(lines)
    finally:
        ole.close()


def _printable_doc_lines(decoded: str) -> list[str]:
    cleaned_chars: list[str] = []
    for char in decoded.replace("\r", "\n").replace("\x07", "\n"):
        codepoint = ord(char)
        if char in "\n\t":
            cleaned_chars.append(char)
        elif 32 <= codepoint <= 126 or 160 <= codepoint <= 591:
            cleaned_chars.append(char)
        else:
            cleaned_chars.append("\n")
    text = "".join(cleaned_chars)
    lines: list[str] = []
    for raw_line in re.split(r"[\n]+", text):
        line = re.sub(r"\s+", " ", raw_line).strip()
        if len(line) < 2:
            continue
        if not any(char.isalpha() for char in line):
            continue
        lines.append(line)
    return lines


def _looks_like_binary_noise(line: str) -> bool:
    if not line:
        return True
    alpha_num = sum(char.isalnum() for char in line)
    if alpha_num / max(len(line), 1) < 0.35:
        return True
    suspicious = sum(1 for char in line if ord(char) > 383)
    return suspicious > max(3, len(line) // 4)


def _extract_eml_text(path: Path) -> tuple[str, list[Path]]:
    with path.open("rb") as handle:
        message = BytesParser(policy=policy.default).parse(handle)
    parts: list[str] = []
    attachments: list[Path] = []
    attachment_dir = _attachment_output_dir(path)
    for header in ("Subject", "From", "To", "Date"):
        value = message.get(header)
        if value:
            parts.append(f"{header}: {value}")
    if parts:
        parts.append("")
    body_parts: list[str] = []
    html_parts: list[str] = []
    for part in message.walk() if message.is_multipart() else [message]:
        if str(part.get_content_disposition() or "").lower() == "attachment":
            filename = part.get_filename()
            if filename:
                safe_name = _safe_attachment_name(filename, f"attachment-{len(attachments) + 1}")
                target = _unique_attachment_path(attachment_dir, safe_name)
                payload = part.get_payload(decode=True)
                if payload is None:
                    try:
                        content = part.get_content()
                    except Exception:
                        content = ""
                    payload = (
                        content.encode(part.get_content_charset() or "utf-8", errors="ignore")
                        if isinstance(content, str)
                        else b""
                    )
                if payload:
                    target.write_bytes(payload)
                    attachments.append(target)
            continue
        try:
            payload = part.get_content()
        except Exception:
            payload = part.get_payload(decode=True)
            if isinstance(payload, bytes):
                payload = payload.decode(part.get_content_charset() or "utf-8", errors="ignore")
        if not isinstance(payload, str):
            continue
        if part.get_content_type() == "text/html":
            html_parts.append(payload)
        else:
            body_parts.append(payload)
    body = "\n\n".join(part.strip() for part in body_parts if part.strip())
    if not body and html_parts:
        body = "\n\n".join(_strip_html_tags(part) for part in html_parts if part.strip())
    if body:
        parts.append(body.strip())
    if attachments:
        parts.append("")
        parts.append("Attachments: " + ", ".join(item.name for item in attachments))
    return re.sub(r"\n{3,}", "\n\n", "\n".join(parts)).strip(), attachments


def _extract_msg_text(path: Path) -> tuple[str, list[Path]]:
    if extract_msg is None:
        raise RuntimeError("extract-msg is not installed.")
    msg = None
    try:
        msg = extract_msg.Message(path)

        def coerce(value: object) -> str:
            if value is None:
                return ""
            if isinstance(value, bytes):
                for encoding in ("utf-8", "cp1252", "latin-1"):
                    try:
                        return value.decode(encoding).replace("\x00", "").strip()
                    except UnicodeDecodeError:
                        continue
                return value.decode("utf-8", errors="ignore").replace("\x00", "").strip()
            return str(value).replace("\x00", "").strip()

        parts: list[str] = []
        for label, attr in (("Subject", "subject"), ("From", "sender"), ("To", "to"), ("Cc", "cc"), ("Date", "date")):
            try:
                value = coerce(getattr(msg, attr, None))
            except Exception:
                value = ""
            if value:
                parts.append(f"{label}: {value}")
        if parts:
            parts.append("")
        body = coerce(getattr(msg, "body", None))
        if body.lstrip()[:100].lower().startswith(("<html", "<body", "<!doctype")):
            body = _strip_html_tags(body)
        if not body:
            body = _strip_html_tags(coerce(getattr(msg, "htmlBody", None)))
        if body:
            parts.append(body.strip())
        attachments: list[Path] = []
        attachment_dir = _attachment_output_dir(path)
        for attachment in list(getattr(msg, "attachments", []) or []):
            name = ""
            for attr in ("longFilename", "shortFilename", "displayName"):
                try:
                    name = coerce(getattr(attachment, attr, None))
                except Exception:
                    name = ""
                if name:
                    break
            if name:
                safe_name = _safe_attachment_name(name, f"attachment-{len(attachments) + 1}")
                target = _unique_attachment_path(attachment_dir, safe_name)
                data = None
                for attr in ("data", "content"):
                    try:
                        candidate = getattr(attachment, attr, None)
                    except Exception:
                        candidate = None
                    if isinstance(candidate, bytes):
                        data = candidate
                        break
                if data:
                    target.write_bytes(data)
                    attachments.append(target)
                    continue
                before = {item.name for item in attachment_dir.iterdir() if item.is_file()}
                try:
                    attachment.save(customPath=str(attachment_dir))
                except TypeError:
                    try:
                        attachment.save()
                    except Exception:
                        continue
                except Exception:
                    continue
                after = [item for item in attachment_dir.iterdir() if item.is_file() and item.name not in before]
                if after:
                    attachments.extend(after)
        if attachments:
            parts.append("")
            parts.append("Attachments: " + ", ".join(str(item.name) for item in attachments))
        return re.sub(r"\n{3,}", "\n\n", "\n".join(parts)).strip(), attachments
    finally:
        if msg is not None:
            try:
                msg.close()
            except Exception:
                pass


def _image_record(source_file: SourceFile) -> ImageRecord:
    width = height = None
    if PILImage is not None:
        try:
            with PILImage.open(source_file.path) as image:
                width, height = image.size
        except Exception:
            width = height = None
    return ImageRecord(
        image_id=f"image-{source_file.sha256[:16]}",
        source_file_id=source_file.file_id,
        path=source_file.path,
        sha256=source_file.sha256,
        width=width,
        height=height,
    )


def _normalized_companion_text(path: Path) -> str | None:
    root = Path("docs/reference/normalized")
    if not root.exists():
        return None
    source_name = path.name.lower()
    for companion in root.glob("collisionrelateddocs__instructions__*.md"):
        if source_name not in companion.name.lower().replace("-", " "):
            continue
        text = companion.read_text(encoding="utf-8", errors="replace")
        if "Conversion status: blocked or incomplete" in text:
            continue
        marker = "---"
        parts = text.split(marker, 2)
        return parts[-1].strip() if len(parts) >= 3 else text
    return None


def read_document(source_file: SourceFile) -> DocumentModel:
    path = source_file.path
    suffix = path.suffix.lower()
    warnings = list(source_file.warnings)
    try:
        if suffix in {".txt", ".md", ".json"}:
            return DocumentModel(source_file, _read_text_file(path), "instruction", "plain_text", warnings)
        if suffix == ".pdf":
            for reader in (_extract_pdf_pymupdf, _extract_pdf_pdfplumber, _extract_pdf_pypdf):
                try:
                    text, notes, method = reader(path)
                    images = _pdf_page_image_records(path, source_file) if source_file.role_guess == "image_pack" else []
                    return DocumentModel(source_file, text, source_file.role_guess, method, warnings + notes, images=images)
                except Exception:
                    continue
            images = _pdf_page_image_records(path, source_file)
            if images:
                warnings.append(
                    ParserWarning(
                        code="image_only_pdf_requires_review",
                        message="PDF contains page images but no native text; OCR or manual review is required.",
                    )
                )
                return DocumentModel(
                    source_file,
                    "",
                    source_file.role_guess,
                    "pdf_image_metadata",
                    warnings,
                    images=images,
                )
            raise RuntimeError("No PDF reader could extract text.")
        if suffix == ".docx":
            return DocumentModel(source_file, _extract_docx_text(path), source_file.role_guess, "python_docx_ooxml", warnings)
        if suffix == ".doc":
            for method_name, reader in (
                ("rtf_plain_text", _extract_rtf_text),
                ("word_com", _extract_doc_text_via_word),
                ("libreoffice_docx", _extract_doc_text_via_soffice),
                ("antiword", _extract_doc_text_via_antiword),
                ("olefile_worddocument", _extract_doc_text_via_olefile),
            ):
                try:
                    return DocumentModel(source_file, reader(path), source_file.role_guess, method_name, warnings)
                except Exception:
                    continue
            companion = _normalized_companion_text(path)
            if companion:
                warnings.append(
                    ParserWarning(
                        code="normalized_companion_fallback",
                        message="Legacy DOC extraction used the normalized companion because native conversion failed.",
                    )
                )
                return DocumentModel(source_file, companion, source_file.role_guess, "normalized_companion", warnings)
            raise RuntimeError("No DOC reader could extract text.")
        if suffix == ".eml":
            text, attachments = _extract_eml_text(path)
            return DocumentModel(source_file, text, "email", "email_parser", warnings, attachments=attachments)
        if suffix == ".msg":
            text, attachments = _extract_msg_text(path)
            return DocumentModel(source_file, text, "email", "extract_msg", warnings, attachments=attachments)
        if suffix in {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tif", ".tiff", ".webp", ".heic"}:
            return DocumentModel(
                source_file,
                "",
                "image",
                "image_metadata",
                warnings,
                images=[_image_record(source_file)],
            )
    except Exception as exc:
        companion = _normalized_companion_text(path)
        if companion:
            warnings.append(
                ParserWarning(
                    code="normalized_companion_fallback",
                    message=f"Native extraction failed ({exc}); used normalized companion.",
                )
            )
            return DocumentModel(source_file, companion, source_file.role_guess, "normalized_companion", warnings)
        warnings.append(ParserWarning(code="reader_failed", message=str(exc), severity="blocker"))
        return DocumentModel(source_file, "", "unknown", "failed", warnings)
    warnings.append(ParserWarning(code="unsupported_file_type", message=f"{suffix} requires manual review."))
    return DocumentModel(source_file, "", "unknown", "unsupported", warnings)


def image_records_for_paths(paths: list[Path], source_file_id: str) -> list[ImageRecord]:
    records: list[ImageRecord] = []
    for path in paths:
        if not path.exists() or not path.is_file():
            continue
        sha = file_sha256(path)
        source = SourceFile(
            file_id=f"file-{sha[:16]}",
            path=path,
            sha256=sha,
            extension=path.suffix.lower(),
            media_type="image/*",
            role_guess="evidence_image",
        )
        records.append(_image_record(source))
    return records

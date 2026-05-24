from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import textwrap
import zipfile
from collections import Counter, defaultdict
from datetime import date, datetime
from pathlib import Path
from typing import Any

try:
    import openpyxl
except Exception:  # pragma: no cover - reported in generated docs
    openpyxl = None

try:
    from docx import Document
except Exception:  # pragma: no cover - reported in generated docs
    Document = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - reported in generated docs
    PdfReader = None


ROOT = Path(__file__).resolve().parents[1]
TODAY = date.today().isoformat()

RAW_SOURCE_ROOT = ("docs", "reference", "raw", "collisionrelateddocs")
REFERENCE_DATA_ROOT = ("docs", "reference", "data")
NORMALIZED_ROOT = ("docs", "reference", "normalized")
ARCHIVED_PLANS_ROOT = ("docs", "plans", "operational-core", "archived_plans")

GENERATED_REFERENCE_PREFIXES = {
    "originalplans_output",
    "ce_system_plans_enhanced",
    "ce_phase4_agents_reviewed_plan",
    "phase7_expanded_markdown_plan",
    "collision_engineers_ai_tools_plans_markdown",
    "collision_engineers_bulk_data_research_pack",
    "cedocumentmapper_rebuild_plan_pack_all_zips",
}

TEST_CONTEXT_PREFIXES = {"testprojectcontext"}

GENERATED_PACK_FOLDERS = [
    "originalplans_output",
    "ce_system_plans_enhanced",
    "ce_phase4_agents_reviewed_plan",
    "phase7_expanded_markdown_plan",
    "collision_engineers_ai_tools_plans_markdown",
    "collision_engineers_bulk_data_research_pack",
    "cedocumentmapper_rebuild_plan_pack_all_zips",
]

IGNORED_MANIFEST_PARTS = {
    ".git",
    ".obsidian",
    ".pytest_cache",
    "__pycache__",
    ".mypy_cache",
    ".ruff_cache",
    ".venv",
    "venv",
    "env",
    "outputs",
    "tmp",
    "dist",
    "build",
    "node_modules",
}
IGNORED_MANIFEST_NAMES = {".DS_Store", "Thumbs.db"}
IGNORED_MANIFEST_SUFFIXES = {".pyc", ".pyo", ".pyd", ".log", ".tmp", ".bak"}

TEXT_EXTENSIONS = {".md", ".txt", ".json"}
RAW_DOC_EXTENSIONS = {".pdf", ".docx", ".doc", ".DOC", ".msg", ".xlsm", ".xlsx", ".jam", ".zip"}


def rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(textwrap.dedent(content).strip() + "\n", encoding="utf-8")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def slugify(value: str) -> str:
    value = value.replace("\\", "/")
    value = re.sub(r"[^A-Za-z0-9._/-]+", "-", value)
    value = value.replace("/", "__")
    return value.strip("-").lower()


def classify(path: Path) -> str:
    parts = path.relative_to(ROOT).parts
    first = parts[0] if parts else ""
    if parts[: len(RAW_SOURCE_ROOT)] == RAW_SOURCE_ROOT:
        return "raw-source"
    if parts[: len(NORMALIZED_ROOT)] == NORMALIZED_ROOT:
        return "normalized-derivative"
    if parts[: len(REFERENCE_DATA_ROOT)] == REFERENCE_DATA_ROOT:
        return "normalized-derivative"
    if len(parts) >= 3 and parts[0] == "docs" and parts[1] == "reference" and parts[2] in {"originalplanning", "test-context"}:
        return "generated-reference"
    if len(parts) >= 4 and parts[0] == "docs" and parts[1] == "plans" and parts[3] == "archived_plans":
        return "implemented" if "implemented" in parts else "superseded"
    if first in GENERATED_REFERENCE_PREFIXES:
        return "generated-reference"
    if first in TEST_CONTEXT_PREFIXES:
        return "generated-reference"
    if first in {"docs", "src", "tests", "tools"}:
        return "active-plan" if path.suffix.lower() in {".md", ".py", ".json", ".toml"} else "normalized-derivative"
    if path.name == "initrepoplan.md":
        return "implemented"
    if path.name in {"README.md", "AGENTS.md", ".gitignore", ".gitattributes", "pyproject.toml"}:
        return "active-plan"
    return "parked"


def source_group(path: Path) -> str:
    parts = path.relative_to(ROOT).parts
    if not parts:
        return "unknown"
    if parts[: len(RAW_SOURCE_ROOT)] == RAW_SOURCE_ROOT:
        if len(parts) > len(RAW_SOURCE_ROOT):
            return f"raw-evidence/{parts[len(RAW_SOURCE_ROOT)]}"
        return "raw-evidence"
    if parts[: len(NORMALIZED_ROOT)] == NORMALIZED_ROOT:
        return "normalized"
    if parts[: len(REFERENCE_DATA_ROOT)] == REFERENCE_DATA_ROOT:
        if len(parts) > len(REFERENCE_DATA_ROOT):
            return f"data/{parts[len(REFERENCE_DATA_ROOT)]}"
        return "data"
    if len(parts) >= 4 and parts[0] == "docs" and parts[1] == "plans" and parts[3] == "archived_plans":
        return "archived_plans"
    if len(parts) >= 4 and parts[0] == "docs" and parts[1] == "reference" and parts[2] == "originalplanning":
        return f"originalplanning/{parts[3]}"
    if len(parts) >= 4 and parts[0] == "docs" and parts[1] == "reference" and parts[2] == "test-context":
        return f"test-context/{parts[3]}"
    return parts[0]


def extraction_plan(path: Path) -> tuple[str, str, str]:
    ext = path.suffix.lower()
    path_str = rel(path)
    if path.name == "source_manifest.csv":
        return "generated-csv", "complete", ""
    if ext == ".md":
        return "native-markdown", "complete", ""
    if ext == ".txt":
        return "native-text", "complete", ""
    if ext == ".json":
        return "json-copy", "complete", ""
    if ext == ".pdf":
        if PdfReader is None:
            return "blocked", "blocked", "pypdf is not available in the active runtime"
        return "pypdf-text", "partial-review-required", "No OCR/layout validation in scaffold pass"
    if ext == ".docx":
        if Document is None:
            return "blocked", "blocked", "python-docx is not available in the active runtime"
        return "python-docx-paragraphs-tables", "partial-review-required", "Headers, footers, text boxes, and embedded objects require later OOXML/manual QA"
    if ext == ".doc":
        return "blocked", "blocked", "Legacy .DOC needs Microsoft Word automation or LibreOffice/soffice; neither is configured"
    if ext == ".msg":
        return "blocked", "blocked", "MSG extraction needs extract-msg or equivalent parser; not installed"
    if ext in {".xlsx", ".xlsm"}:
        if openpyxl is None:
            return "blocked", "blocked", "openpyxl is not available"
        return "openpyxl-workbook-summary", "partial-review-required", "Workbook formulas/macros/images are not executed or visually validated"
    if ext == ".zip":
        return "zip-entry-list", "partial-review-required", "ZIP contents are inventoried but not recursively promoted"
    if ext == ".jam":
        if path_str == "docs/reference/raw/collisionrelateddocs/collision_releated/Collision Engineers Whiteboard.jam":
            return "jam-zip-image-export + Figma get_figjam", "partial-review-required", ""
        return "blocked", "blocked", "Google Jamboard format needs manual export/review"
    return "not-applicable", "not-applicable", ""


def companion_path_for(path: Path) -> str:
    if not rel(path).startswith("docs/reference/raw/collisionrelateddocs/"):
        return ""
    raw_relative = rel(path).removeprefix("docs/reference/raw/")
    return f"docs/reference/normalized/{slugify(raw_relative)}.md"


def extract_pdf(path: Path) -> tuple[str, str]:
    if PdfReader is None:
        return "", "pypdf unavailable"
    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            parts.append(f"## Page {index}\n\n{text.strip()}")
        text = "\n\n".join(parts).strip()
        if not text:
            return "", "No extractable text found; OCR/manual QA required"
        return text, ""
    except Exception as exc:
        return "", f"PDF extraction failed: {exc}"


def extract_docx(path: Path) -> tuple[str, str]:
    if Document is None:
        return "", "python-docx unavailable"
    try:
        doc = Document(str(path))
        lines: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        for table_index, table in enumerate(doc.tables, start=1):
            lines.append(f"\n## Table {table_index}")
            for row in table.rows:
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        text = "\n\n".join(lines).strip()
        if not text:
            return "", "No paragraph/table text found; manual QA required"
        return text, ""
    except Exception as exc:
        return "", f"DOCX extraction failed: {exc}"


def workbook_summary(path: Path, max_rows: int = 25) -> tuple[str, str]:
    if openpyxl is None:
        return "", "openpyxl unavailable"
    try:
        wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
        lines = [f"# Workbook Companion: {path.name}", "", f"- Sheets: {', '.join(wb.sheetnames)}", ""]
        for ws in wb.worksheets:
            lines.append(f"## Sheet: {ws.title}")
            lines.append(f"- Dimensions: {ws.max_row} rows x {ws.max_column} columns")
            lines.append("")
            rows = []
            for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, max_rows), values_only=True):
                rows.append([format_cell(value) for value in row[: min(ws.max_column, 12)]])
            if rows:
                width = len(rows[0])
                lines.append("| " + " | ".join([f"Column {i}" for i in range(1, width + 1)]) + " |")
                lines.append("| " + " | ".join(["---"] * width) + " |")
                for row in rows:
                    row = row + [""] * (width - len(row))
                    lines.append("| " + " | ".join(escape_md(cell) for cell in row) + " |")
            lines.append("")
        return "\n".join(lines), ""
    except Exception as exc:
        return "", f"Workbook summary failed: {exc}"


def format_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    return str(value)


def escape_md(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ").strip()


def write_companion(path: Path, record: dict[str, str]) -> None:
    companion_rel = record["normalized_companion"]
    if not companion_rel:
        return
    companion = ROOT / companion_rel
    ext = path.suffix.lower()
    body = ""
    blocker = record["blocker"]
    if ext == ".pdf":
        body, blocker = extract_pdf(path)
    elif ext == ".docx":
        body, blocker = extract_docx(path)
    elif ext in {".xlsx", ".xlsm"}:
        body, blocker = workbook_summary(path)
    elif ext == ".json":
        body = "```json\n" + path.read_text(encoding="utf-8", errors="replace") + "\n```"
    elif ext == ".txt":
        body = "```text\n" + path.read_text(encoding="utf-8", errors="replace") + "\n```"
    elif ext == ".md":
        body = path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".zip":
        try:
            with zipfile.ZipFile(path) as zf:
                entries = sorted(zf.namelist())
            body = "\n".join(f"- `{entry}`" for entry in entries)
        except Exception as exc:
            blocker = f"ZIP listing failed: {exc}"
    if not body:
        body = f"Conversion status: blocked or incomplete.\n\nBlocker: {blocker or record['blocker'] or 'Manual review required.'}"
    metadata = textwrap.dedent(
        f"""
        # Normalized Companion: {path.name}

        - Raw source: `{rel(path)}`
        - SHA256: `{record['sha256']}`
        - Extraction method: {record['extraction_method']}
        - Extraction confidence: {record['extraction_confidence']}
        - Blocker: {blocker or record['blocker'] or 'None recorded'}

        This companion is a derivative working copy. The raw source remains the source of truth.

        ---
        """
    ).strip()
    write_text(companion, metadata + "\n\n" + body)


def is_ignored_manifest_artifact(path: Path) -> bool:
    parts = path.relative_to(ROOT).parts
    if set(parts) & IGNORED_MANIFEST_PARTS:
        return True
    if any(part.endswith(".egg-info") for part in parts):
        return True
    name = path.name
    if name == ".env.example":
        return False
    if name == ".env" or name.startswith(".env."):
        return True
    if name in IGNORED_MANIFEST_NAMES:
        return True
    return any(name.endswith(suffix) for suffix in IGNORED_MANIFEST_SUFFIXES)


def iter_files() -> list[Path]:
    files = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        if is_ignored_manifest_artifact(path):
            continue
        files.append(path)
    return sorted(files, key=lambda p: rel(p).lower())


def build_manifest() -> list[dict[str, str]]:
    records = []
    for path in iter_files():
        if rel(path).startswith("docs/reference/normalized/"):
            status = "normalized-derivative"
            method = "generated-companion"
            confidence = "derived"
            blocker = ""
            companion = ""
        else:
            status = classify(path)
            method, confidence, blocker = extraction_plan(path)
            companion = companion_path_for(path)
        records.append(
            {
                "path": rel(path),
                "status": status,
                "source_group": source_group(path),
                "file_type": path.suffix.lower() or "[no extension]",
                "size_bytes": str(path.stat().st_size),
                "sha256": sha256(path),
                "canonical_path": rel(path),
                "normalized_companion": companion,
                "extraction_method": method,
                "extraction_confidence": confidence,
                "blocker": blocker,
            }
        )
    return records


def safe_move_directory(source: Path, target: Path) -> None:
    if not source.exists():
        return
    source_resolved = source.resolve()
    target_parent = target.parent.resolve()
    root_resolved = ROOT.resolve()
    if not str(source_resolved).lower().startswith(str(root_resolved).lower()):
        raise RuntimeError(f"Refusing to move source outside workspace: {source_resolved}")
    if not str(target_parent).lower().startswith(str(root_resolved).lower()):
        raise RuntimeError(f"Refusing to move target outside workspace: {target_parent}")
    if target.exists():
        if not source.is_dir() or not target.is_dir():
            raise RuntimeError(f"Refusing to overwrite existing path during move: {target}")
        for child in sorted(source.iterdir(), key=lambda item: item.name.lower()):
            destination = target / child.name
            if destination.exists():
                if child.is_dir() and destination.is_dir():
                    safe_move_directory(child, destination)
                else:
                    raise RuntimeError(f"Refusing to overwrite existing path during merge: {destination}")
            else:
                shutil.move(str(child), str(destination))
        try:
            source.rmdir()
        except OSError as exc:
            raise RuntimeError(f"Could not remove merged source directory {source}: {exc}") from exc
    else:
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(source), str(target))


def remove_empty_directory(path: Path) -> None:
    if not path.exists():
        return
    resolved = path.resolve()
    root_resolved = ROOT.resolve()
    if not str(resolved).lower().startswith(str(root_resolved).lower()):
        raise RuntimeError(f"Refusing to remove directory outside workspace: {resolved}")
    if not path.is_dir():
        raise RuntimeError(f"Expected directory for cleanup: {path}")
    try:
        path.rmdir()
    except OSError:
        return


def organize_reference_packs() -> None:
    safe_move_directory(ROOT / "collisionrelateddocs", ROOT / "docs/reference/raw/collisionrelateddocs")
    safe_move_directory(ROOT / "docs/normalized", ROOT / "docs/reference/normalized")
    safe_move_directory(ROOT / "docs/data", ROOT / "docs/reference/data")
    safe_move_directory(ROOT / "archive/plans", ROOT / "docs/plans/operational-core/archived_plans")
    remove_empty_directory(ROOT / "archive")
    safe_move_directory(ROOT / "docs/planning", ROOT / "docs/plans/operational-core")
    safe_move_directory(ROOT / "docs/tickets", ROOT / "docs/plans/operational-core/tickets")
    for folder in GENERATED_PACK_FOLDERS:
        safe_move_directory(ROOT / folder, ROOT / "docs/reference/originalplanning" / folder)
    safe_move_directory(ROOT / "docs/reference/generated-packs", ROOT / "docs/reference/originalplanning")
    safe_move_directory(ROOT / "testprojectcontext", ROOT / "docs/reference/test-context/testprojectcontext")


def write_generated_pack_index() -> None:
    rows = []
    for base in [ROOT / "docs/reference/originalplanning", ROOT / "docs/reference/test-context"]:
        if not base.exists():
            continue
        for folder in sorted([item for item in base.iterdir() if item.is_dir()], key=lambda p: p.name.lower()):
            file_count = sum(1 for item in folder.rglob("*") if item.is_file())
            rows.append((folder.name, rel(folder), file_count))
    lines = [
        "# Original Planning Index",
        "",
        "Generated and historical planning packs are reference-only unless promoted into an active ticket or canonical doc.",
        "",
        "| Pack | Path | Files | Status |",
        "| --- | --- | ---: | --- |",
    ]
    for name, path, file_count in rows:
        lines.append(f"| `{escape_md(name)}` | `{escape_md(path)}` | {file_count} | generated-reference |")
    write_text(ROOT / "docs/reference/originalplanning_index.md", "\n".join(lines))


def write_manifest(records: list[dict[str, str]]) -> None:
    docs = ROOT / "docs"
    docs.mkdir(exist_ok=True)
    fieldnames = list(records[0].keys()) if records else []
    with (docs / "source_manifest.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(records)
    write_text(docs / "source_manifest.json", json.dumps(records, indent=2))

    status_counts = Counter(r["status"] for r in records)
    ext_counts = Counter(r["file_type"] for r in records)
    group_counts = Counter(r["source_group"] for r in records)
    lines = [
        "# Source Manifest",
        "",
        f"- Generated: {TODAY}",
        f"- Total files inventoried: {len(records)}",
        "- Full machine-readable inventory: `docs/source_manifest.csv` and `docs/source_manifest.json`",
        "",
        "## Status Counts",
        "",
        "| Status | Count |",
        "| --- | ---: |",
    ]
    for status, count in sorted(status_counts.items()):
        lines.append(f"| {status} | {count} |")
    lines.extend(["", "## Source Group Counts", "", "| Source group | Count |", "| --- | ---: |"])
    for group, count in sorted(group_counts.items()):
        lines.append(f"| {group} | {count} |")
    lines.extend(["", "## Extension Counts", "", "| Extension | Count |", "| --- | ---: |"])
    for ext, count in sorted(ext_counts.items()):
        lines.append(f"| {ext} | {count} |")
    lines.extend(
        [
            "",
            "## Inventory",
            "",
            "| Path | Status | Type | Size | SHA256 | Normalized companion | Blocker |",
            "| --- | --- | --- | ---: | --- | --- | --- |",
        ]
    )
    for r in records:
        lines.append(
            "| "
            + " | ".join(
                [
                    f"`{escape_md(r['path'])}`",
                    r["status"],
                    r["file_type"],
                    r["size_bytes"],
                    f"`{r['sha256']}`",
                    f"`{escape_md(r['normalized_companion'])}`" if r["normalized_companion"] else "",
                    escape_md(r["blocker"]),
                ]
            )
            + " |"
        )
    write_text(docs / "source_manifest.md", "\n".join(lines))


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def normalize_code(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip().upper()


def provider_config_rows() -> list[dict[str, str]]:
    path = ROOT / "docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json"
    data = load_json(path)
    rows = []
    for provider in data["providers"]:
        work_provider = provider.get("field_rules", {}).get("work_provider", {})
        code = normalize_code(work_provider.get("config"))
        rows.append(
            {
                "provider_name": provider.get("name", ""),
                "parser_code": code,
                "detect_phrases": "; ".join(provider.get("detect_phrases", [])),
                "engineer_report": str(provider.get("engineer_report", False)),
            }
        )
    return rows


def read_principals_from_workbook(path: Path) -> list[dict[str, str]]:
    if openpyxl is None:
        return []
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    ws = wb["Principals"]
    rows = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        name, eva_code, box_code, inbox, instructions, drag, sent_mino = row[1:8]
        if not any([name, eva_code, box_code, inbox, instructions, drag, sent_mino]):
            continue
        rows.append(
            {
                "source": rel(path),
                "principal_name": format_cell(name),
                "eva_code": normalize_code(eva_code),
                "box_code": normalize_code(box_code),
                "inbox": format_cell(inbox),
                "instructions": format_cell(instructions),
                "drag_into_eva": format_cell(drag),
                "sent_mino": format_cell(sent_mino),
            }
        )
    return rows


def read_actual_jobs(path: Path) -> Counter[str]:
    if openpyxl is None:
        return Counter()
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    ws = wb["Jobs"]
    values = Counter()
    for row in ws.iter_rows(min_row=7, values_only=True):
        principal = normalize_code(row[4] if len(row) > 4 else None)
        if principal:
            values[principal] += 1
    return values


def read_mapped_principals(path: Path) -> tuple[list[str], dict[str, str]]:
    if openpyxl is None:
        return [], {}
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    ws = wb.active
    principals = []
    lost = {}
    for row in ws.iter_rows(min_row=3, values_only=True):
        principal = normalize_code(row[1] if len(row) > 1 else None)
        lost_code = normalize_code(row[3] if len(row) > 3 else None)
        lost_reason = format_cell(row[4] if len(row) > 4 else None)
        if principal:
            principals.append(principal)
        if lost_code:
            lost[lost_code] = lost_reason
    return principals, lost


def write_provider_matrix() -> None:
    output = ROOT / "docs/reference/data/provider_coverage_matrix.csv"
    output.parent.mkdir(parents=True, exist_ok=True)
    parser_rows = provider_config_rows()
    parser_names = {normalize_code(row["provider_name"]) for row in parser_rows}
    parser_codes = {row["parser_code"] for row in parser_rows if row["parser_code"]}
    covered_codes = parser_names | parser_codes

    wb_paths = [
        ROOT / "docs/reference/raw/collisionrelateddocs/collision_releated/Backup of CE Job Sheet 260429.xlsm",
        ROOT / "docs/reference/raw/collisionrelateddocs/collision_releated/Backup of CE Job Sheet 260309.xlsm",
    ]
    principal_by_code: dict[str, dict[str, str]] = {}
    for wb_path in wb_paths:
        for row in read_principals_from_workbook(wb_path):
            code = row["eva_code"] or row["box_code"] or normalize_code(row["principal_name"])
            if code and code not in principal_by_code:
                principal_by_code[code] = row

    actual_jobs = Counter()
    for wb_path in wb_paths:
        actual_jobs.update(read_actual_jobs(wb_path))

    mapped, lost = read_mapped_principals(ROOT / "docs/reference/raw/collisionrelateddocs/collision_releated/Mapped Principals.xlsx")
    all_codes = set(covered_codes) | set(principal_by_code) | set(actual_jobs) | set(mapped) | set(lost)
    rows = []
    for code in sorted(c for c in all_codes if c):
        p = principal_by_code.get(code, {})
        parser_matches = [
            row["provider_name"]
            for row in parser_rows
            if normalize_code(row["provider_name"]) == code or row["parser_code"] == code
        ]
        rows.append(
            {
                "code": code,
                "principal_name": p.get("principal_name", ""),
                "parser_covered": "yes" if parser_matches else "no",
                "parser_provider_names": "; ".join(parser_matches),
                "job_sheet_principals_table": "yes" if code in principal_by_code else "no",
                "actual_jobs_count": str(actual_jobs.get(code, 0)),
                "mapped_principals_entry": "yes" if code in mapped else "no",
                "lost_cause": "yes" if code in lost else "no",
                "lost_cause_reason": lost.get(code, ""),
                "eva_code": p.get("eva_code", ""),
                "box_code": p.get("box_code", ""),
                "inbox": p.get("inbox", ""),
                "instructions": p.get("instructions", ""),
                "drag_into_eva": p.get("drag_into_eva", ""),
            }
        )
    with output.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)

    uncovered_actual = [r for r in rows if r["actual_jobs_count"] != "0" and r["parser_covered"] == "no"]
    uncovered_table = [r for r in rows if r["job_sheet_principals_table"] == "yes" and r["parser_covered"] == "no"]
    mapped_only = [
        r
        for r in rows
        if r["mapped_principals_entry"] == "yes"
        and r["job_sheet_principals_table"] == "no"
        and r["parser_covered"] == "no"
    ]

    md = [
        "# Provider Coverage Matrix",
        "",
        f"- Generated: {TODAY}",
        "- Parser source: `docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json`",
        "- Job-sheet sources: `Backup of CE Job Sheet 260429.xlsm`, `Backup of CE Job Sheet 260309.xlsm`",
        "- Mapped-principals source: `Mapped Principals.xlsx`",
        "- Full matrix: `docs/reference/data/provider_coverage_matrix.csv`",
        "",
        "## Summary",
        "",
        f"- Parser provider presets: {len(parser_rows)}",
        f"- Unique parser coverage codes/names: {len(covered_codes)}",
        f"- Job-sheet principal table codes: {len(principal_by_code)}",
        f"- Actual job principal codes observed: {len(actual_jobs)}",
        f"- Actual job principal codes not parser-covered: {', '.join(r['code'] for r in uncovered_actual) or 'None'}",
        f"- Job-sheet principal table codes not parser-covered: {len(uncovered_table)}",
        f"- Mapped-only uncovered codes: {len(mapped_only)}",
        "",
        "## Parser Presets",
        "",
        "| Provider preset | Parser code | Engineer report | Detect phrases |",
        "| --- | --- | --- | --- |",
    ]
    for row in parser_rows:
        md.append(
            f"| `{escape_md(row['provider_name'])}` | `{escape_md(row['parser_code'])}` | {row['engineer_report']} | {escape_md(row['detect_phrases'])} |"
        )
    md.extend(["", "## Actual Job Principals Not Parser-Covered", "", "| Code | Count | Lost cause | Reason |", "| --- | ---: | --- | --- |"])
    for row in uncovered_actual:
        md.append(f"| `{row['code']}` | {row['actual_jobs_count']} | {row['lost_cause']} | {escape_md(row['lost_cause_reason'])} |")
    md.extend(
        [
            "",
            "## Job-Sheet Principal Table Codes Not Parser-Covered",
            "",
            "| Code | Principal name | Inbox | Instructions | Drag into EVA |",
            "| --- | --- | --- | --- | --- |",
        ]
    )
    for row in uncovered_table:
        md.append(
            f"| `{row['code']}` | {escape_md(row['principal_name'])} | {escape_md(row['inbox'])} | {escape_md(row['instructions'])} | {escape_md(row['drag_into_eva'])} |"
        )
    md.extend(["", "## Mapped-Only Uncovered Codes", "", ", ".join(f"`{r['code']}`" for r in mapped_only) or "None"])
    write_text(ROOT / "docs/reference/data/provider_coverage_matrix.md", "\n".join(md))


def write_spreadsheet_flow_doc() -> None:
    content = """
    # CE Job Sheet Spreadsheet Companion

    ## Sources

    - `docs/reference/raw/collisionrelateddocs/collision_releated/Backup of CE Job Sheet 260429.xlsm`
    - `docs/reference/raw/collisionrelateddocs/collision_releated/Backup of CE Job Sheet 260309.xlsm`
    - `docs/reference/raw/collisionrelateddocs/collision_releated/Mapped Principals.xlsx`

    ## Observed Workbook Structure

    Both CE Job Sheet backups contain a `Jobs` sheet, `Principals` sheet, and `Garages` sheet. The newer backup also contains `Own figures`.

    ## Operational Reading

    The `Jobs` sheet is the operational queue. Rows near the top include cases missing something or waiting for instructions/images; rows lower in the sheet are jobs ready for EVA once images, instructions, timing, and other blockers are resolved.

    The `Principals` sheet is the provider/principal operational lookup. It maps solicitor or work-provider names to EVA codes, Box codes, inbox ownership, instruction receipt style, whether a job can be dragged into EVA, and notes such as direct-job handling or WhatsApp dependence.

    The `Garages` sheet is an inspection-address/contact lookup. It contains garage name, address, email, phone, and figure-related notes. This is operationally sensitive and must be treated as raw source data.

    `Mapped Principals.xlsx` appears to be a coverage/control sheet linking current parser provider names to known principals, plus lost-cause entries such as `ACSP` where OCR quality is too low.

    ## Automation Implications

    - Do not treat the parser provider config as complete coverage for operational providers.
    - Build a provider coverage matrix before adding new parser rules.
    - Separate parser preset coverage from EVA/Box principal-code coverage.
    - Preserve raw spreadsheet files unchanged. Generated summaries and CSVs are derivative and must link back to source hashes.
    - Do not execute workbook macros during source analysis. Macro behavior must be documented separately before any automation depends on it.

    ## Known Spreadsheet Flow Requirements

    - Missing instruction/image cases must remain visible as blocked/manual-review states.
    - Jobs can be blocked if images are absent, instructions are absent, images cannot be linked to a case, or the requested date is in the future.
    - Garage contact method varies by garage and can include email or WhatsApp outside the workbook.
    - Principal code generation resets yearly and must be implemented as a controlled reference generator, not a spreadsheet side effect.
    """
    write_text(ROOT / "docs/operations/job_sheet_spreadsheet_companion.md", textwrap.dedent(content))


def write_static_docs() -> None:
    write_text(
        ROOT / "README.md",
        """
        # Collision Command Centre

        Collision Command Centre (CCC) is the canonical private repository for rebuilding Collision Engineers' vehicle-damage instruction parser and related evidence-handling foundations.

        ## Current Scope

        The first programme milestone is the Operational Core MVP. The parser is the first executable MVP inside it:

        - shared Python parser core;
        - non-technical office-staff UI;
        - equivalent CLI for automation and AI-agent usage;
        - provider/principal admin;
        - work-item and review queue;
        - provider configuration and coverage tracking;
        - EVA-ready JSON/payload validation;
        - Box-ready evidence-package generation;
        - raw-source preservation with normalized Markdown/metadata companions.

        Collision Engineers do not do personal injury or KADOE work. Those workflows must not be planned into CCC.

        ## Source Of Truth

        Raw operational files under `docs/reference/raw/collisionrelateddocs/` remain immutable source evidence. Generated plan packs are reference-only unless promoted into canonical docs or tickets. The current canonical entry points are:

        - `docs/docs_index.md`
        - `docs/repo_map.json`
        - `docs/source_manifest.md`
        - `docs/roadmap.md`
        - `docs/plans/_index.md`
        - `docs/plans/operational-core/source_synthesis.md`
        - `docs/plans/parser-extraction/parser-mvp/plan.md`
        - `docs/architecture/`
        - `docs/contracts/`
        - `docs/decisions/`
        - `docs/plans/operational-core/tickets/backlog_index.md`
        - `docs/requirements/`
        - `docs/reference/data/provider_coverage_matrix.md`
        - `docs/reference/originalplanning_index.md`
        - `docs/plans/*/archived_plans/implemented/`

        ## Local Verification

        Run:

        ```powershell
        $py='C:\\Users\\PC\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe'
        & $py tools/verify_scaffold.py
        ```

        The verifier checks the promoted planning scaffold, provider coverage evidence, decision docs, parser MVP plan coverage, and source-manifest integrity.
        """,
    )
    write_text(
        ROOT / "AGENTS.md",
        """
        # AGENTS.md

        ## Navigation

        - Start with `docs/docs_index.md` for human-readable navigation and `docs/repo_map.json` for machine-readable path routing.
        - Initial repository setup and exhaustive reference-derived idea planning live under `docs/plans/initial-repo-setup/`.
        - Active programme planning lives under `docs/plans/operational-core/`.
        - Current parser MVP implementation work lives at `docs/plans/parser-extraction/parser-mvp/plan.md`; `docs/plans/operational-core/parser-mvp/plan.md` is only a compatibility stub.
        - Active tickets live under `docs/plans/operational-core/tickets/`.
        - Implemented or superseded plans live under the owning workspace's `archived_plans/`.
        - Raw evidence lives under `docs/reference/raw/collisionrelateddocs/`.
        - Normalized companions and extracted data live under `docs/reference/normalized/` and `docs/reference/data/`.
        - Generated or historical planning packs live under `docs/reference/originalplanning/` and are reference-only unless promoted.

        ## Repository Hygiene

        - Treat `docs/reference/raw/collisionrelateddocs/` raw files as immutable evidence. Do not edit, rename destructively, overwrite, or normalize in place.
        - Keep generated derivatives under `docs/reference/normalized/`, `docs/reference/data/`, or another documented derivative path.
        - Update `docs/source_manifest.md`, `docs/source_manifest.csv`, and `docs/source_manifest.json` whenever source files, generated companions, active docs, or archives change.
        - Keep `README.md`, `docs/roadmap.md`, `docs/architecture/`, `docs/contracts/`, `docs/requirements/`, and `docs/glossary.md` current with behavior-changing work.
        - Do not commit secrets, API keys, tokens, OAuth material, mailbox credentials, or provider credentials. Run the scaffold verification and a targeted secret scan before any commit or push.
        - Preserve unrelated local configuration, auth state, MCP settings, and user files.
        - Use `rg` or `rg --files` first for search. Use structured parsers for JSON, spreadsheets, DOCX, and PDFs when available.
        - Use `apply_patch` for manual edits. Do not create or edit files with shell write tricks.

        ## Documentation Sync Rules

        - At the start of any task, check `docs/docs_index.md`, `docs/repo_map.json`, `docs/roadmap.md`, the owning workspace `plan.md`, active tickets, relevant ADRs/contracts, and source evidence before changing files.
        - At the start of parser work, also check `docs/plans/parser-extraction/parser-mvp/plan.md`, `docs/plans/parser-extraction/parser-mvp/adjacent-parser-and-inspection-location-review.md`, `docs/reference/adjacent_repositories.md`, and the provider/corpus sources named by the parser plan.
        - `docs/roadmap.md` must show both the full programme plan and the current project position. Update its Current Status section whenever the active phase, milestone, blockers, or next action changes.
        - On completion of any task, update the owning plan/ticket, `docs/roadmap.md`, `docs/docs_index.md`, `docs/repo_map.json`, and any affected architecture, contract, decision, requirement, operations, security, or glossary docs.
        - If docs, generated companions, active plans, tickets, archives, or source inventories change, regenerate `docs/source_manifest.md`, `docs/source_manifest.csv`, and `docs/source_manifest.json`.
        - A task that changes behavior, scope, status, roadmap position, source ownership, or verification requirements is not complete until the relevant documentation and manifests are updated or the omission is explicitly documented.

        ## Planning And Ticket Lifecycle

        - Every active plan or ticket must include status, owner, created date, last reviewed date, source links, roadmap milestone, dependencies, expected outputs, acceptance criteria, verification required, archive target, and supersedes/superseded-by fields.
        - When a plan or ticket is implemented, move it to the owning workspace's `archived_plans/implemented/`, rename it with the completion date, and add an implemented-state block at the top.
        - Superseded or merged plans go to the owning workspace's `archived_plans/superseded/` with a pointer to the replacement doc or ticket.
        - Generated plan packs are reference material unless explicitly promoted into active docs or tickets.

        ## Parser Rules

        - Keep UI and CLI thin. Both must call the same parser core and share validation/export contracts.
        - Do not import the legacy `cedocumentmapper` monolith wholesale. Reference it for behavior and migrate requirements deliberately.
        - Parser output must pass canonical schema validation before EVA-specific output.
        - Provider coverage must be tracked in `docs/reference/data/provider_coverage_matrix.md` before adding or changing provider rules.
        - Any cloud OCR/document-intelligence path requires privacy, cost, data-residency, and vendor review before use.

        ## Verification

        - Verify before claiming completion. Prefer reproducible commands and direct checks.
        - For this scaffold, run `tools/verify_scaffold.py`.
        - If pytest or external parser dependencies are not installed, document that instead of claiming test coverage from unavailable tools.
        """,
    )
    write_text(
        ROOT / "docs/docs_index.md",
        """
        # Documentation Index

        This is the primary human navigation file for CCC documentation.

        ## Quick Start

        - Product scope and current milestone: `README.md` and `docs/roadmap.md`.
        - Initial setup planning: `docs/plans/initial-repo-setup/README.md`.
        - Active programme plan: `docs/plans/operational-core/source_synthesis.md`.
        - Planned folder taxonomy: `docs/plans/initial-repo-setup/documentation-scaffold/plans-folder-expansion-plan.md`.
        - Local repo task-start skill plan and install record: `docs/plans/initial-repo-setup/local-repo-task-start-skill-plan.md`.
        - Parser MVP plan: `docs/plans/parser-extraction/parser-mvp/plan.md`.
        - Active backlog: `docs/plans/operational-core/tickets/backlog_index.md`.
        - Agent path map: `docs/repo_map.json`.
        - Full source inventory: `docs/source_manifest.md`, `docs/source_manifest.csv`, and `docs/source_manifest.json`.

        ## Active Work

        | Area | Path | Use |
        | --- | --- | --- |
        | Plans | `docs/plans/_index.md` | Active plan workspaces and archived plans. |
        | Initial Repo Setup | `docs/plans/initial-repo-setup/` | Pre-code repository setup, documentation scaffold, and exhaustive reference idea planning. |
        | Operational Core | `docs/plans/operational-core/` | Current programme plan, parser MVP plan, and tickets. |
        | Architecture | `docs/architecture/` | System architecture and programme boundaries. |
        | Contracts | `docs/contracts/` | Versioned schemas and integration contracts. |
        | Decisions | `docs/decisions/` | ADRs and option papers. |
        | Requirements | `docs/requirements/` | Business requirements and open questions. |
        | Operations | `docs/operations/` | Runbooks, monitoring, release, rollback, and spreadsheet companion notes. |
        | Security | `docs/security/` | Data map, vendor register, DPIA, safety review, and API security. |

        ## Reference Material

        | Area | Path | Use |
        | --- | --- | --- |
        | Raw evidence | `docs/reference/raw/collisionrelateddocs/` | Immutable source files. Do not edit in place. |
        | Normalized companions | `docs/reference/normalized/` | Generated Markdown companions for raw evidence. |
        | Reference data | `docs/reference/data/` | Provider matrix and extracted Jam/FigJam derivatives. |
        | Original planning | `docs/reference/originalplanning/` | Historical/generated planning packs; reference-only unless promoted. |
        | Test context | `docs/reference/test-context/` | Historical test repositories and context packs. |

        ## Quality Rules

        - At task start, read the roadmap, repo map, owning workspace plan, active tickets, and relevant source evidence before changing files.
        - At completion of any task, update `docs/roadmap.md`, the owning plan/ticket, `docs/docs_index.md`, `docs/repo_map.json`, affected key docs, and `docs/source_manifest.*`.
        - Update `docs/source_manifest.*` when source files, generated companions, active docs, or archives change.
        - Promote ideas from reference material into `docs/plans/initial-repo-setup/reference-audit/`, `docs/plans/operational-core/tickets/`, or another owning plan before treating them as active scope.
        - Keep raw evidence immutable and create derivatives under `docs/reference/normalized/` or `docs/reference/data/`.
        - Run `python tools/verify_scaffold.py` after documentation structure changes.
        """,
    )
    write_text(
        ROOT / "docs/plans/_index.md",
        """
        # Plans Index

        `docs/plans/` contains active plan workspaces and plan archives.

        | Plan Area | Path | Status |
        | --- | --- | --- |
        | Initial Repo Setup | `docs/plans/initial-repo-setup/` | active pre-code planning workspace |
        | Operational Core | `docs/plans/operational-core/` | active programme workspace |
        | Workspace Expansion Plan | `docs/plans/initial-repo-setup/documentation-scaffold/plans-folder-expansion-plan.md` | active taxonomy plan for future plan folders |
        | Repository Restructure | `docs/plans/initial-repo-setup/archived_plans/implemented/2026-05-23-implemented-repository-restructure.md` | implemented |

        ## Initial Repo Setup Layout

        - `docs/plans/initial-repo-setup/documentation-scaffold/` captures repository setup and documentation scaffold planning.
        - `docs/plans/initial-repo-setup/reference-audit/` captures exhaustive reference-derived idea inventories.
        - `docs/plans/initial-repo-setup/tickets/` contains pre-code setup tickets.
        - `docs/plans/initial-repo-setup/archived_plans/` contains implemented or superseded setup plans.

        ## Operational Core Layout

        - `docs/plans/operational-core/source_synthesis.md` maps reference material into canonical outputs.
        - `docs/plans/parser-extraction/parser-mvp/plan.md` is the current parser MVP implementation plan.
        - `docs/plans/operational-core/tickets/` contains active phased tickets.
        - `docs/plans/operational-core/archived_plans/implemented/` contains completed plans.
        - `docs/plans/operational-core/archived_plans/superseded/` contains superseded or merged plans.

        ## Planned Expansion Workspaces

        The saved expansion plan assigns reference-derived work to these future folders before implementation starts:

        | Planned Area | Future Path | Owns |
        | --- | --- | --- |
        | Unified Platform | `docs/plans/unified-platform/` | Mature end-to-end platform and convergence roadmap. |
        | Automation Centre | `docs/plans/automation-centre/` | Deterministic automation model, workflows, retries, exception routing, and automation KPIs. |
        | Parser Extraction | `docs/plans/parser-extraction/` | Parser, extraction adapters, mapper compatibility, OCR fallback, and corpus regression. |
        | Case Workflow State | `docs/plans/case-workflow-state/` | Work item state, review queue, audit stream, missing-info state, duplicate and historical search. |
        | Provider Principal Config | `docs/plans/provider-principal-config/` | Provider, principal, garage, routing, and provider-admin planning. |
        | Intake Storage Integrations | `docs/plans/intake-storage-integrations/` | Outlook, Box, EVA/Sentry, website/WhatsApp metadata, and spreadsheet bridge adapters. |
        | Evidence Estimate Review | `docs/plans/evidence-estimate-review/` | Evidence, image quality, estimate parsing, ABP review, duplicate/reused evidence, and damage workbench planning. |
        | Vehicle Valuation Data | `docs/plans/vehicle-valuation-data/` | Vehicle intelligence, DVLA/DVSA, MOT data, valuation evidence, salvage, mileage, and identity planning. |
        | Engineer Communications | `docs/plans/engineer-communications/` | Engineer pack, templates, report drafting support, status summaries, and communication workflows. |
        | AI Agents | `docs/plans/ai-agents/` | Controlled workflow agents and agent-vs-automation boundaries. |
        | MCP And Tooling | `docs/plans/mcp-and-tooling/` | MCP servers, internal tool adapters, registry, schemas, audit, rate limits, and security gateway. |
        | Agent Skills | `docs/plans/agent-skills/` | Reusable drafting, summary, RAG, valuation, communication, and training skills. |
        | AI Platform Tools | `docs/plans/ai-platform-tools/` | Model evaluation, prompt governance, datasets, AI feedback loops, and policy implementation. |
        | User Experience Interfaces | `docs/plans/user-experience-interfaces/` | Staff, engineer, admin, parser, review, dashboard, portal/front-door, and accessibility planning. |
        | Finance Billing | `docs/plans/finance-billing/` | Invoice, fee-note, payment status, chaser metadata, approvals, and payment automation option papers. |
        | Governance Security | `docs/plans/governance-security/` | DPIA, vendor governance, privacy, PII redaction, data licensing, expert-boundary, and API security planning. |
        | Operations Quality | `docs/plans/operations-quality/` | Test corpus, regression harness, release/rollback, monitoring, runbooks, rollout, and decommissioning. |
        | Analytics Data Platform | `docs/plans/analytics-data-platform/` | Operations analytics, data warehouse, EVA mining, data quality, risk indicators, scheduling, and BI. |
        | External Platform Partners | `docs/plans/external-platform-partners/` | Customer portal, partner API, insurer integrations, Audatex partnerships, and partner access controls. |
        | Product Business | `docs/plans/product-business/` | Discovery, ROI, client pitch, conservative positioning, objection handling, and defensibility. |
        """,
    )
    write_text(
        ROOT / "docs/reference/_index.md",
        """
        # Reference Index

        `docs/reference/` contains source evidence, generated derivatives, historical planning, and test-context material.

        | Area | Path | Status |
        | --- | --- | --- |
        | Raw evidence | `docs/reference/raw/collisionrelateddocs/` | immutable source evidence |
        | Normalized companions | `docs/reference/normalized/` | generated derivatives |
        | Reference data | `docs/reference/data/` | generated/reference datasets and exports |
        | Original planning | `docs/reference/originalplanning/` | generated historical planning packs |
        | Test context | `docs/reference/test-context/` | historical test repositories/context packs |

        Generated and historical material is not active scope unless promoted into `docs/plans/initial-repo-setup/reference-audit/`, architecture, contracts, decisions, or `docs/plans/operational-core/tickets/`.
        """,
    )
    write_text(
        ROOT / "docs/repo_map.json",
        json.dumps(
            {
                "version": 1,
                "updated": TODAY,
                "entry_points": {
                    "human_index": "docs/docs_index.md",
                    "machine_map": "docs/repo_map.json",
                    "source_manifest": "docs/source_manifest.md",
                    "roadmap": "docs/roadmap.md",
                },
                "plans": {
                    "index": "docs/plans/_index.md",
                    "workspace_expansion_plan": "docs/plans/initial-repo-setup/documentation-scaffold/plans-folder-expansion-plan.md",
                    "initial_repo_setup": {
                        "root": "docs/plans/initial-repo-setup",
                        "readme": "docs/plans/initial-repo-setup/README.md",
                        "documentation_scaffold": "docs/plans/initial-repo-setup/documentation-scaffold",
                        "local_repo_task_start_skill_plan": "docs/plans/initial-repo-setup/local-repo-task-start-skill-plan.md",
                        "reference_audit": "docs/plans/initial-repo-setup/reference-audit",
                        "tickets": "docs/plans/initial-repo-setup/tickets",
                        "implemented_archive": "docs/plans/initial-repo-setup/archived_plans/implemented",
                        "superseded_archive": "docs/plans/initial-repo-setup/archived_plans/superseded",
                    },
                    "operational_core": {
                        "root": "docs/plans/operational-core",
                        "source_synthesis": "docs/plans/operational-core/source_synthesis.md",
                        "parser_mvp_plan": "docs/plans/parser-extraction/parser-mvp/plan.md",
                        "tickets": "docs/plans/operational-core/tickets",
                        "implemented_archive": "docs/plans/operational-core/archived_plans/implemented",
                        "superseded_archive": "docs/plans/operational-core/archived_plans/superseded",
                    },
                    "planned_workspaces": {
                        "unified_platform": "docs/plans/unified-platform",
                        "automation_centre": "docs/plans/automation-centre",
                        "parser_extraction": "docs/plans/parser-extraction",
                        "case_workflow_state": "docs/plans/case-workflow-state",
                        "provider_principal_config": "docs/plans/provider-principal-config",
                        "intake_storage_integrations": "docs/plans/intake-storage-integrations",
                        "evidence_estimate_review": "docs/plans/evidence-estimate-review",
                        "vehicle_valuation_data": "docs/plans/vehicle-valuation-data",
                        "engineer_communications": "docs/plans/engineer-communications",
                        "ai_agents": "docs/plans/ai-agents",
                        "mcp_and_tooling": "docs/plans/mcp-and-tooling",
                        "agent_skills": "docs/plans/agent-skills",
                        "ai_platform_tools": "docs/plans/ai-platform-tools",
                        "user_experience_interfaces": "docs/plans/user-experience-interfaces",
                        "finance_billing": "docs/plans/finance-billing",
                        "governance_security": "docs/plans/governance-security",
                        "operations_quality": "docs/plans/operations-quality",
                        "analytics_data_platform": "docs/plans/analytics-data-platform",
                        "external_platform_partners": "docs/plans/external-platform-partners",
                        "product_business": "docs/plans/product-business",
                    },
                    "repository_restructure": "docs/plans/initial-repo-setup/archived_plans/implemented/2026-05-23-implemented-repository-restructure.md",
                },
                "reference": {
                    "index": "docs/reference/_index.md",
                    "raw_evidence": "docs/reference/raw/collisionrelateddocs",
                    "normalized_companions": "docs/reference/normalized",
                    "reference_data": "docs/reference/data",
                    "provider_coverage_matrix": "docs/reference/data/provider_coverage_matrix.md",
                    "original_planning": "docs/reference/originalplanning",
                    "original_planning_index": "docs/reference/originalplanning_index.md",
                    "test_context": "docs/reference/test-context",
                },
                "active_docs": {
                    "architecture": "docs/architecture",
                    "contracts": "docs/contracts",
                    "decisions": "docs/decisions",
                    "requirements": "docs/requirements",
                    "operations": "docs/operations",
                    "security": "docs/security",
                    "glossary": "docs/glossary.md",
                },
                "rules": {
                    "raw_evidence": "Do not edit files under docs/reference/raw/collisionrelateddocs in place.",
                    "reference_material": "Reference material is not active scope unless promoted into plans, tickets, architecture, contracts, or decisions.",
                    "manifest": "Update docs/source_manifest.md, docs/source_manifest.csv, and docs/source_manifest.json whenever source files, generated companions, active docs, or archives change.",
                },
            },
            indent=2,
        ),
    )
    write_text(
        ROOT / ".gitignore",
        """
        __pycache__/
        .pytest_cache/
        .mypy_cache/
        .ruff_cache/
        .venv/
        venv/
        env/
        *.pyc
        *.pyo
        *.pyd
        *.log
        *.tmp
        *.bak
        .env
        .env.*
        !.env.example
        outputs/
        tmp/
        dist/
        build/
        *.egg-info/
        node_modules/
        .DS_Store
        Thumbs.db
        """,
    )
    write_text(
        ROOT / ".gitattributes",
        """
        * text=auto
        *.md text eol=lf
        *.py text eol=lf
        *.json text eol=lf
        *.csv text eol=lf
        *.pdf binary
        *.doc binary
        *.DOC binary
        *.docx binary
        *.msg binary
        *.xlsm binary
        *.xlsx binary
        *.zip binary
        *.jam binary
        """,
    )
    write_text(
        ROOT / "docs/glossary.md",
        """
        # Glossary

        | Term | Meaning |
        | --- | --- |
        | CCC | Collision Command Centre. |
        | CE | Collision Engineers. |
        | EVA | System used to enter and manage cases after instructions and images are ready. |
        | Sentry | EVA/API-adjacent system described in the supplied docs. |
        | Principal | Internal work-provider code used by Collision Engineers and EVA. |
        | Case/PO | Collision Engineers internal reference. It is provider-coded and increments yearly. |
        | Claim no | Client/work-provider claim reference. |
        | VRM | Vehicle registration mark. |
        | Instruction | Client/work-provider document requesting inspection or assessment. |
        | Image-based assessment | Assessment using vehicle images rather than a physical inspection address. |
        | Companion report | Valuation or vehicle-check evidence report downloaded from a supporting service. |
        | Box | Initial evidence storage and case package destination. |
        | Raw source | Immutable original source file. |
        | Normalized companion | Generated derivative Markdown/text/metadata file linked to a raw source hash. |
        """,
    )
    write_text(
        ROOT / "docs/roadmap.md",
        """
        # Roadmap

        ## Current Status

        Current position: Section 0 complete for parser readiness. The next implementation phase is Section 2 parser MVP execution under `docs/plans/parser-extraction/parser-mvp/plan.md`.

        Current milestone: begin parser MVP implementation from the committed documentation baseline, using the parser MVP plan, adjacent parser review, provider sources, and local repo task-start skill.

        Parser implementation status: ready to start; parser code work has not started in this phase. The active parser MVP plan is `docs/plans/parser-extraction/parser-mvp/plan.md`; the operational-core parser path is only a compatibility stub.

        Pre-parser readiness gates: repository documentation lifecycle rules are explicit, this Current Status section is up to date, source manifests match the working tree, scaffold verification and scaffold contract tests pass, and parser work should start from a committed/pushed documentation baseline.

        Local repo skill status: installed at `C:/Users/PC/.codex/skills/ccc-repo-task-start` and planned/evaluated in `docs/plans/initial-repo-setup/local-repo-task-start-skill-plan.md`. This local development skill is separate from `docs/plans/agent-skills/`, which is reserved for production Collision Engineers skills.

        ## P0 - Repository Safety And Documentation Foundation

        1. Initialize repository hygiene files.
        2. Inventory all files in `docs/source_manifest.md`, `.csv`, and `.json`.
        3. Preserve raw operational files unchanged.
        4. Generate normalized companions or blocker notes for raw sources.
        5. Classify generated plan packs as reference-only unless promoted.

        ## P1 - Parser-First Foundation

        1. Build the shared parser core under `src/ccc_parser/`.
        2. Expose equivalent UI and CLI surfaces over the same parser core.
        3. Load current provider presets from `docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json`.
        4. Add canonical parser result models and EVA export contract boundaries.
        5. Add verification scripts and tests.

        ## P2 - Provider Coverage And Golden Tests

        1. Maintain `docs/reference/data/provider_coverage_matrix.md`.
        2. Prioritize the 26 current parser presets.
        3. Separate live job-sheet uncovered principals from historical/mapped-only codes.
        4. Add golden fixtures and expected outputs for each provider.

        ## P3 - Evidence Storage And Integrations

        1. Define Box case-package contract.
        2. Keep storage behind an adapter boundary.
        3. Prepare future path to CCC-owned storage on Google Cloud, AWS, or Azure.
        4. Review privacy, data residency, cost, and vendor obligations before cloud OCR or document intelligence.

        ## P4 - Workflow Expansion

        Only start broader workflow automation after parser reliability, UI/CLI parity, canonical schema, provider coverage, and archive/ticket lifecycle are proven.
        """,
    )
    write_text(
        ROOT / "docs/architecture/overview.md",
        """
        # Architecture Overview

        CCC is parser-first. The initial architecture separates source preservation, extraction, provider rules, canonical data, EVA export, UI, CLI, and storage adapters.

        ```mermaid
        flowchart LR
            Raw["Raw instructions and evidence"] --> Companion["Normalized companions and metadata"]
            Raw --> Readers["Document readers"]
            Readers --> Core["Parser core"]
            Provider["Provider config"] --> Core
            Core --> Canonical["Canonical parser result"]
            Canonical --> Validation["Validation and warnings"]
            Validation --> UI["Office-staff UI"]
            Validation --> CLI["Automation CLI"]
            Validation --> EVA["EVA JSON/payload export"]
            Validation --> Package["Case evidence package"]
            Package --> Box["Box adapter first"]
            Package --> Future["Future cloud storage adapter"]
        ```

        ## Boundaries

        - UI and CLI do not implement extraction rules. They call the parser core.
        - Provider detection and field extraction are configuration-backed and testable.
        - EVA export is an adapter from canonical parser output, not the canonical model itself.
        - Box is the first storage integration, but storage must remain replaceable.
        - Cloud OCR/document intelligence is a later adapter, not the default parser runtime.
        """,
    )
    write_text(
        ROOT / "docs/architecture/parser_ui_cli.md",
        """
        # Parser UI And CLI Architecture

        ## UI Requirement

        The initial parser is not fully automated. Non-technical office staff must be able to import files, inspect extraction results, correct fields, view warnings, export EVA-ready JSON/payloads, and package images/evidence without using a terminal.

        Required UI capabilities:

        - drag/drop and file-picker import for PDFs, DOCX, DOC, MSG/EML, images, and batches;
        - batch navigation and blocked/manual-review queues;
        - source text preview;
        - extracted field review and correction;
        - provider preset selection and provider mismatch warning;
        - image extraction/export area;
        - EVA JSON/payload preview;
        - validation warnings for missing VRM, provider, images, mileage, or inspection address;
        - audit metadata and blocker report export.

        ## CLI Requirement

        The CLI must expose the same parser core behavior for automation and AI-agent usage:

        - import/extract;
        - batch parse;
        - validate;
        - export JSON/payload;
        - export images;
        - emit audit metadata and blocker sidecars.

        ## Technology Decision

        The scaffold keeps the UI implementation as a thin Python module. The detailed UI technology decision remains open between modernized Tkinter, a Windows desktop UI, and a local web UI. The default bias is the lowest-friction Windows office deployment that does not block a later dashboard.
        """,
    )
    write_text(
        ROOT / "docs/architecture/tooling.md",
        """
        # Local Tooling Requirements

        ## Current Availability

        Available in the bundled Python runtime:

        - `pypdf`
        - `python-docx`
        - `openpyxl`
        - `pandas`
        - PIL

        Not currently available on PATH or in the bundled Python runtime:

        - LibreOffice/`soffice`
        - Pandoc
        - Tesseract
        - PyMuPDF
        - pdfplumber
        - `extract-msg`
        - OCRmyPDF
        - MarkItDown
        - Outlook/Word COM libraries

        ## Recommended Local-First Stack

        - PDF native text: PyMuPDF primary, `pypdf` fallback.
        - PDF tables/layout: `pdfplumber` targeted fallback.
        - DOCX: `python-docx` plus direct OOXML inspection.
        - Legacy `.DOC`: Microsoft Word automation where Office is installed, with LibreOffice headless conversion as a service/CI fallback.
        - MSG: `extract-msg` or equivalent pure-Python parser.
        - XLSM/XLSX: `openpyxl`/`pandas` for structure and values. Do not execute macros.
        - OCR: Tesseract/pytesseract as local fallback; OCRmyPDF only as optional batch normalization.
        - Markdown normalization: MarkItDown can help bulk companions, but CCC parser extraction must remain provenance-aware and rule-backed.

        LibreOffice is important only for local headless conversion of legacy Office files. It is not part of Collision Engineers' business workflow.
        """,
    )
    write_text(
        ROOT / "docs/contracts/parser_result_v1.md",
        """
        # Parser Result Contract v1

        Parser output is canonical before it becomes EVA-specific.

        ## Required Concepts

        - source file path and hash;
        - detected provider and provider confidence;
        - extracted fields with value, confidence, source span/page where available, and warning state;
        - validation warnings;
        - image extraction/package metadata;
        - audit metadata for tool version and extraction methods.

        ## Core Fields

        - work provider/principal;
        - VRM;
        - vehicle model;
        - claimant/client name;
        - client claim/reference number;
        - incident date;
        - instruction date;
        - inspection date;
        - inspection address or image-based assessment marker;
        - accident circumstances;
        - VAT status;
        - mileage and mileage unit;
        - evidence image list.

        ## Rule

        EVA field shape is an export adapter. Do not let EVA-specific JSON become the only internal representation.
        """,
    )
    write_text(
        ROOT / "docs/contracts/eva_export_contract.md",
        """
        # EVA Export Contract

        EVA-ready means local validation and manual review, not direct Sentry submission.

        ## Export Requirements

        - Preserve the field order and shape represented by `docs/reference/raw/collisionrelateddocs/Final Format Example 02.json`.
        - Export only after required parser warnings are reviewed.
        - Include source audit metadata outside the EVA payload when EVA does not support it.
        - Keep image ordering requirements separate from field JSON: first two preview images, then all images including the first two again.

        ## Open Items

        - Confirm final EVA JSON field names against live EVA import behavior.
        - Add golden tests for all 26 parser presets.
        - Decide how manual corrections are represented in audit metadata.
        """,
    )
    write_text(
        ROOT / "docs/contracts/storage_adapter_contract.md",
        """
        # Storage Adapter Contract

        Box is the first storage integration. Future CCC-owned storage may move to Google Cloud, AWS, or Azure.

        ## Adapter Boundary

        Storage adapters receive a validated case package and return storage references. Parser extraction and EVA export must not know provider-specific SDK details.

        ## Case Package Contents

        - original instruction file;
        - original email where available;
        - images;
        - EVA-ready JSON/payload;
        - valuation or companion report where available;
        - audit metadata;
        - blocker/manual-review notes where relevant.

        ## Box First

        - Folder naming follows case/PO number.
        - Metadata must capture principal, claim reference, VRM, source hashes, and upload status.
        - Live upload requires credentials and retention decisions outside this scaffold.

        ## Future Storage

        A later architecture decision can choose Google Cloud, AWS, or Azure. Azure remains a serious candidate because Azure Document Intelligence may align document extraction with storage/hosting, but no cloud processing is allowed without privacy, cost, data-residency, and vendor review.
        """,
    )
    write_text(
        ROOT / "docs/requirements/business_requirements.md",
        """
        # Business Requirements

        ## Scope

        CCC supports vehicle-damage instruction parsing and evidence preparation for Collision Engineers.

        Collision Engineers do not do personal injury or KADOE work. Those workflows are out of scope and must not be planned into CCC.

        ## Current Workflow

        1. Instructions arrive by email in three Outlook mailboxes.
        2. Staff enter details into the CE Job Sheet.
        3. Staff chase vehicle damage images when absent.
        4. Once instructions and images are available, staff enter the case into EVA.
        5. Staff add mileage, valuation evidence, Experian/adverse-history checks, and inspection address where required.
        6. Images must be added in a required order.
        7. Case evidence is stored in Box by case/PO number.
        8. The case is handed to an engineer.

        ## MVP Requirements

        - Preserve current provider parsing behavior before expanding.
        - Provide a full UI for office staff.
        - Provide equivalent CLI functionality for automation.
        - Track missing images/instructions as blockers.
        - Generate validated EVA-ready output.
        - Preserve raw evidence and source traceability.
        """,
    )
    write_text(
        ROOT / "docs/requirements/open_business_questions.md",
        """
        # Open Business Questions

        - Which valuation provider should be the system of record: Glass's, CAP HPI, Auto Trader, Experian/HPI, or another source?
        - What is the exact yearly reset process for case/PO numbering by principal?
        - Which garages use WhatsApp vs email, and who owns each contact route?
        - What is the minimum acceptable image set for each provider?
        - Which providers should be prioritized after the 26 current parser presets?
        - What Box retention, folder permissions, and metadata rules are required?
        - What storage provider should CCC own long-term: Google Cloud, AWS, or Azure?
        - What adverse-history/Experian evidence is required before EVA entry?
        """,
    )
    write_text(
        ROOT / "docs/decisions/0001-parser-first.md",
        """
        # ADR 0001: Parser-First Repository Direction

        - Status: accepted
        - Date: 2026-05-23

        ## Decision

        CCC starts with the parser foundation, office-staff UI, equivalent CLI, source manifest, provider coverage, and export/storage contracts before broader workflow automation.

        ## Rationale

        The parser is already the strongest automation path because EVA can consume JSON imports. The new source documents clarify missing provider coverage and spreadsheet workflows. Rebuilding around a shared parser core reduces drift between office use and future automation.

        ## Consequences

        - Workflow dashboard, Outlook ingestion, WhatsApp automation, valuation automation, and direct Sentry submission are later work.
        - UI and CLI must share the parser core.
        - Provider coverage and golden tests become gating work.
        """,
    )
    write_text(
        ROOT / "docs/decisions/0002-storage-box-first.md",
        """
        # ADR 0002: Box First, Storage Adapter Always

        - Status: accepted
        - Date: 2026-05-23

        ## Decision

        Plan Box as the first evidence-storage integration, but put all storage behavior behind an adapter contract.

        ## Rationale

        Collision Engineers currently categorize case evidence in Box by case/PO number. Future CCC-owned storage may move to Google Cloud, AWS, or Azure, so parser/export logic must not depend on Box SDK details.

        ## Consequences

        - Box metadata and folder conventions are documented early.
        - Future storage selection remains an architecture decision.
        - Cloud OCR/document intelligence requires privacy and vendor review.
        """,
    )
    write_text(
        ROOT / "docs/integrations/external_services.md",
        """
        # External Services And Vendor Review

        ## Box

        Box is the initial storage integration. CCC needs an adapter that can create or locate a case folder, upload evidence files, apply metadata, and return stable storage references. Live upload is not part of this scaffold.

        Official docs: https://developer.box.com/

        ## Azure Document Intelligence

        Azure Document Intelligence is a future candidate if CCC storage/hosting moves to Azure. It can support OCR and document layout extraction, but any use must pass privacy, cost, data-residency, and vendor review.

        Official docs:

        - https://learn.microsoft.com/azure/ai-services/document-intelligence/
        - https://learn.microsoft.com/azure/ai-services/document-intelligence/prebuilt/read

        ## AWS Textract

        AWS Textract is a future candidate for OCR/form/table extraction if CCC chooses AWS.

        Official docs: https://docs.aws.amazon.com/textract/

        ## Google Document AI

        Google Document AI is a future candidate if CCC chooses Google Cloud storage/hosting.

        Official docs: https://cloud.google.com/document-ai/docs

        ## Local-First Rule

        The parser-first MVP uses local extraction/OCR by default. Cloud processing is an adapter decision, not an implicit fallback.
        """,
    )
    write_text(
        ROOT / "docs/security/source_safety_review.md",
        """
        # Source Safety Review

        - Review date: 2026-05-23
        - Scope: scaffold-time text/config scan and corpus risk classification.

        ## Result

        No obvious live plaintext API key, password, token, or private key was identified in the text/config scan performed during scaffold implementation.

        The corpus is still operationally sensitive. Raw Office, PDF, MSG, spreadsheet, and normalized companion files can contain personal data, garage contact details, claim references, vehicle registrations, addresses, phone numbers, and client/provider information.

        ## Scan Boundary

        The scan covered text-readable project files and excluded binary Office/PDF/MSG/archive formats from deep content inspection. Binary files are treated as raw evidence, not sanitized material.

        The scan produced expected documentation hits for words such as `token`, `secret`, `password`, and placeholder examples in Sentry/EVA planning docs. Those are not treated as live credentials without real credential values.

        ## Rules

        - Keep this repository private.
        - Do not push to any remote until remote privacy and access controls are confirmed.
        - Do not commit real credentials.
        - Store future Box, EVA/Sentry, Microsoft, valuation, DVLA/DVSA, AWS, Azure, or Google credentials in a proper secrets store.
        - Redact or separate production evidence before creating any public demo or shareable sample.
        """,
    )
    write_text(
        ROOT / "docs/reference/adjacent_repositories.md",
        """
        # Adjacent Repository Review

        ## `../cedocumentmapper`

        Legacy monolith. It combines Tkinter UI, provider config, document parsing, OCR, MSG handling, and export behavior in one app. Treat it as behavior reference, not code to import wholesale.

        ## `../cedocumentmapper_v2.0`

        Contract-first rewrite scaffold. Useful concepts include domain models, reader protocols, rule protocols, contracts, and tickets. Production app behavior is not implemented.

        ## `../collisionautomation`

        React/Vite prototype with useful UI, matching, schemas, runtime gating, and testing patterns. It is not parser-first.

        ## `../collisionpdf`

        Closest parser-first reference. It has FastAPI parser service modules for native extraction, OCR fallback, classification, matching, field extraction, and schema validation. Its warning still applies: synthetic fixture tests do not prove real-production accuracy.
        """,
    )
    write_text(
        ROOT / "docs/plans/operational-core/tickets/README.md",
        """
        # Tickets

        Active tickets promoted from generated plan packs belong here.

        Every ticket must include:

        - status;
        - owner;
        - created date;
        - last reviewed date;
        - source links;
        - roadmap milestone;
        - acceptance criteria;
        - verification required;
        - supersedes/superseded-by fields.

        Completed tickets move to `docs/plans/operational-core/archived_plans/implemented/` with an implemented-state block.
        """,
    )


def write_parser_scaffold() -> None:
    write_text(
        ROOT / "pyproject.toml",
        """
        [project]
        name = "ccc"
        version = "0.1.0"
        description = "Collision Command Centre parser-first scaffold"
        requires-python = ">=3.11"
        dependencies = []

        [project.scripts]
        ccc-parser = "ccc_parser.cli:main"

        [tool.pytest.ini_options]
        pythonpath = ["src"]
        testpaths = ["tests"]
        """,
    )
    write_text(ROOT / "src/ccc_parser/__init__.py", '__all__ = ["ParserCore"]\n\nfrom .core import ParserCore\n')
    write_text(
        ROOT / "src/ccc_parser/models.py",
        """
        from __future__ import annotations

        from dataclasses import dataclass, field
        from pathlib import Path
        from typing import Any


        @dataclass(frozen=True)
        class ParserWarning:
            code: str
            message: str
            severity: str = "warning"


        @dataclass(frozen=True)
        class ExtractedField:
            name: str
            value: str | None
            confidence: float = 0.0
            source: str | None = None


        @dataclass(frozen=True)
        class ParserResult:
            source_path: Path
            source_sha256: str
            provider_name: str | None
            fields: dict[str, ExtractedField] = field(default_factory=dict)
            warnings: list[ParserWarning] = field(default_factory=list)
            audit: dict[str, Any] = field(default_factory=dict)

            def to_dict(self) -> dict[str, Any]:
                return {
                    "source_path": str(self.source_path),
                    "source_sha256": self.source_sha256,
                    "provider_name": self.provider_name,
                    "fields": {
                        key: {
                            "name": value.name,
                            "value": value.value,
                            "confidence": value.confidence,
                            "source": value.source,
                        }
                        for key, value in self.fields.items()
                    },
                    "warnings": [
                        {"code": item.code, "message": item.message, "severity": item.severity}
                        for item in self.warnings
                    ],
                    "audit": self.audit,
                }
        """,
    )
    write_text(
        ROOT / "src/ccc_parser/providers.py",
        """
        from __future__ import annotations

        import json
        from dataclasses import dataclass
        from pathlib import Path


        @dataclass(frozen=True)
        class ProviderPreset:
            name: str
            code: str
            detect_phrases: tuple[str, ...]
            engineer_report: bool
            field_rules: dict


        def normalize_code(value: object) -> str:
            return "" if value is None else str(value).strip().upper()


        def load_provider_presets(path: Path) -> list[ProviderPreset]:
            data = json.loads(path.read_text(encoding="utf-8"))
            presets: list[ProviderPreset] = []
            for item in data.get("providers", []):
                work_provider = item.get("field_rules", {}).get("work_provider", {})
                presets.append(
                    ProviderPreset(
                        name=item.get("name", ""),
                        code=normalize_code(work_provider.get("config")),
                        detect_phrases=tuple(item.get("detect_phrases", [])),
                        engineer_report=bool(item.get("engineer_report", False)),
                        field_rules=item.get("field_rules", {}),
                    )
                )
            return presets


        def detect_provider(text: str, presets: list[ProviderPreset]) -> ProviderPreset | None:
            lowered = text.lower()
            for preset in presets:
                if any(phrase.lower() in lowered for phrase in preset.detect_phrases):
                    return preset
            return None
        """,
    )
    write_text(
        ROOT / "src/ccc_parser/core.py",
        """
        from __future__ import annotations

        import hashlib
        from pathlib import Path

        from .models import ParserResult, ParserWarning
        from .providers import detect_provider, load_provider_presets


        DEFAULT_PROVIDER_CONFIG = Path("docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json")


        def file_sha256(path: Path) -> str:
            digest = hashlib.sha256()
            with path.open("rb") as handle:
                for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                    digest.update(chunk)
            return digest.hexdigest()


        class ParserCore:
            def __init__(self, provider_config: Path | None = None) -> None:
                self.provider_config = provider_config or DEFAULT_PROVIDER_CONFIG
                self.providers = load_provider_presets(self.provider_config)

            def parse(self, source_path: Path) -> ParserResult:
                source_path = Path(source_path)
                text = self._best_effort_text(source_path)
                provider = detect_provider(text, self.providers)
                warnings: list[ParserWarning] = []
                if provider is None:
                    warnings.append(
                        ParserWarning(
                            code="provider_not_detected",
                            message="No provider preset matched the available source text.",
                        )
                    )
                warnings.append(
                    ParserWarning(
                        code="scaffold_only",
                        message="Field extraction is scaffolded; provider golden rules are not implemented yet.",
                        severity="info",
                    )
                )
                return ParserResult(
                    source_path=source_path,
                    source_sha256=file_sha256(source_path),
                    provider_name=provider.name if provider else None,
                    warnings=warnings,
                    audit={
                        "provider_config": str(self.provider_config),
                        "provider_count": len(self.providers),
                        "text_preview_length": len(text),
                    },
                )

            def _best_effort_text(self, source_path: Path) -> str:
                if source_path.suffix.lower() in {".txt", ".md", ".json"}:
                    return source_path.read_text(encoding="utf-8", errors="replace")
                return ""
        """,
    )
    write_text(
        ROOT / "src/ccc_parser/exporters/__init__.py",
        """
        from .eva import EVA_FIELD_ORDER

        __all__ = ["EVA_FIELD_ORDER"]
        """,
    )
    write_text(
        ROOT / "src/ccc_parser/exporters/eva.py",
        """
        EVA_FIELD_ORDER = [
            "work_provider",
            "vrm",
            "vehicle_model",
            "claimant_name",
            "reference",
            "incident_date",
            "instruction_date",
            "inspection_date",
            "inspection_address",
            "accident_circumstances",
            "vat_status",
            "mileage",
            "mileage_unit",
        ]
        """,
    )
    write_text(
        ROOT / "src/ccc_parser/cli.py",
        """
        from __future__ import annotations

        import argparse
        import json
        from pathlib import Path

        from .core import DEFAULT_PROVIDER_CONFIG, ParserCore
        from .providers import load_provider_presets


        def build_parser() -> argparse.ArgumentParser:
            parser = argparse.ArgumentParser(prog="ccc-parser")
            parser.add_argument("--provider-config", type=Path, default=DEFAULT_PROVIDER_CONFIG)
            sub = parser.add_subparsers(dest="command", required=True)

            sub.add_parser("providers", help="List configured provider presets")

            parse_cmd = sub.add_parser("parse", help="Parse one source file with scaffold extraction")
            parse_cmd.add_argument("source", type=Path)
            parse_cmd.add_argument("--output", type=Path)
            return parser


        def main(argv: list[str] | None = None) -> int:
            args = build_parser().parse_args(argv)
            if args.command == "providers":
                presets = load_provider_presets(args.provider_config)
                for preset in presets:
                    print(f"{preset.name}\\t{preset.code}")
                return 0
            if args.command == "parse":
                core = ParserCore(args.provider_config)
                result = core.parse(args.source).to_dict()
                payload = json.dumps(result, indent=2)
                if args.output:
                    args.output.parent.mkdir(parents=True, exist_ok=True)
                    args.output.write_text(payload + "\\n", encoding="utf-8")
                else:
                    print(payload)
                return 0
            raise AssertionError(f"Unhandled command: {args.command}")


        if __name__ == "__main__":
            raise SystemExit(main())
        """,
    )
    write_text(
        ROOT / "src/ccc_parser/ui/__init__.py",
        """
        from .app import main

        __all__ = ["main"]
        """,
    )
    write_text(
        ROOT / "src/ccc_parser/ui/app.py",
        """
        from __future__ import annotations

        from pathlib import Path

        from ccc_parser.core import ParserCore


        class ParserUiController:
            def __init__(self, core: ParserCore | None = None) -> None:
                self.core = core or ParserCore()

            def parse_file(self, path: Path):
                return self.core.parse(path)


        def main() -> int:
            import tkinter as tk
            from tkinter import filedialog, messagebox, scrolledtext

            controller = ParserUiController()
            root = tk.Tk()
            root.title("Collision Command Centre Parser")
            root.geometry("980x640")

            output = scrolledtext.ScrolledText(root, wrap=tk.WORD)
            output.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)

            def open_file() -> None:
                selected = filedialog.askopenfilename(
                    title="Select instruction file",
                    filetypes=[
                        ("Supported documents", "*.pdf *.docx *.doc *.msg *.eml *.txt *.md *.json"),
                        ("All files", "*.*"),
                    ],
                )
                if not selected:
                    return
                try:
                    result = controller.parse_file(Path(selected)).to_dict()
                except Exception as exc:
                    messagebox.showerror("Parse failed", str(exc))
                    return
                output.delete("1.0", tk.END)
                output.insert(tk.END, str(result))

            toolbar = tk.Frame(root)
            toolbar.pack(fill=tk.X, padx=12, pady=(12, 0))
            tk.Button(toolbar, text="Import instruction", command=open_file).pack(side=tk.LEFT)
            tk.Label(toolbar, text="Scaffold UI: review/edit/export screens will be built over the shared parser core.").pack(
                side=tk.LEFT, padx=12
            )

            root.mainloop()
            return 0


        if __name__ == "__main__":
            raise SystemExit(main())
        """,
    )
    write_text(
        ROOT / "tests/test_scaffold_contracts.py",
        """
        from pathlib import Path
        import sys

        ROOT = Path(__file__).resolve().parents[1]
        sys.path.insert(0, str(ROOT))

        from ccc_parser.core import ParserCore
        from ccc_parser.exporters import EVA_FIELD_ORDER
        from ccc_parser.providers import load_provider_presets
        from ccc_parser.ui.app import ParserUiController
        import tools.scaffold_initial_repo as scaffold

        def test_provider_config_loads_current_presets():
            presets = load_provider_presets(ROOT / "docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json")
            assert len(presets) == 26
            names = {preset.name for preset in presets}
            assert "ALISON" in names
            assert "TEN" in names


        def test_parser_core_default_provider_config_loads_current_presets(monkeypatch):
            monkeypatch.chdir(ROOT)
            core = ParserCore()
            assert len(core.providers) == 26


        def test_ui_controller_and_cli_share_parser_core():
            core = ParserCore(ROOT / "docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json")
            controller = ParserUiController(core)
            assert controller.core is core


        def test_eva_export_field_order_contains_required_core_fields():
            assert EVA_FIELD_ORDER[:4] == ["work_provider", "vrm", "vehicle_model", "claimant_name"]
            assert "inspection_address" in EVA_FIELD_ORDER
            assert "mileage" in EVA_FIELD_ORDER


        def test_scaffold_manifest_ignores_local_artifacts(monkeypatch, tmp_path):
            monkeypatch.setattr(scaffold, "ROOT", tmp_path)
            (tmp_path / "README.md").write_text("# Test\\n", encoding="utf-8")
            (tmp_path / ".env").write_text("SECRET=value\\n", encoding="utf-8")
            (tmp_path / ".env.example").write_text("SECRET=\\n", encoding="utf-8")
            (tmp_path / ".venv").mkdir()
            (tmp_path / ".venv" / "artifact.py").write_text("ignored\\n", encoding="utf-8")
            (tmp_path / "tmp").mkdir()
            (tmp_path / "tmp" / "scratch.txt").write_text("ignored\\n", encoding="utf-8")
            (tmp_path / "debug.log").write_text("ignored\\n", encoding="utf-8")

            paths = {record["path"] for record in scaffold.build_manifest()}

            assert "README.md" in paths
            assert ".env.example" in paths
            assert ".env" not in paths
            assert ".venv/artifact.py" not in paths
            assert "tmp/scratch.txt" not in paths
            assert "debug.log" not in paths


        def test_scaffold_archive_migration_removes_empty_legacy_root(monkeypatch, tmp_path):
            monkeypatch.setattr(scaffold, "ROOT", tmp_path)
            legacy_plan = tmp_path / "archive" / "plans" / "implemented" / "done.md"
            legacy_plan.parent.mkdir(parents=True)
            legacy_plan.write_text("# Done\\n", encoding="utf-8")

            scaffold.organize_reference_packs()

            migrated = tmp_path / "docs" / "plans" / "operational-core" / "archived_plans" / "implemented" / "done.md"
            assert migrated.exists()
            assert not (tmp_path / "archive").exists()
        """,
    )
    write_text(
        ROOT / "tools/verify_scaffold.py",
        """
        from __future__ import annotations

        import csv
        import json
        import subprocess
        import sys
        from pathlib import Path

        ROOT = Path(__file__).resolve().parents[1]
        sys.path.insert(0, str(ROOT / "src"))

        from ccc_parser.core import ParserCore
        from ccc_parser.providers import load_provider_presets


        REQUIRED_PATHS = [
            "README.md",
            "AGENTS.md",
            ".gitignore",
            ".gitattributes",
            "pyproject.toml",
            "docs/docs_index.md",
            "docs/repo_map.json",
            "docs/source_manifest.md",
            "docs/source_manifest.csv",
            "docs/roadmap.md",
            "docs/plans/_index.md",
            "docs/plans/initial-repo-setup/README.md",
            "docs/plans/initial-repo-setup/local-repo-task-start-skill-plan.md",
            "docs/plans/initial-repo-setup/documentation-scaffold/plan.md",
            "docs/plans/initial-repo-setup/documentation-scaffold/plans-folder-expansion-plan.md",
            "docs/plans/initial-repo-setup/reference-audit/all-ideas-plan.md",
            "docs/plans/initial-repo-setup/tickets/README.md",
            "docs/plans/initial-repo-setup/archived_plans/implemented/.gitkeep",
            "docs/plans/initial-repo-setup/archived_plans/implemented/2026-05-23-implemented-initrepoplan.md",
            "docs/plans/initial-repo-setup/archived_plans/implemented/2026-05-23-implemented-repository-restructure.md",
            "docs/plans/initial-repo-setup/archived_plans/superseded/.gitkeep",
            "docs/architecture/overview.md",
            "docs/architecture/programme_architecture.md",
            "docs/architecture/mvp_interlock.md",
            "docs/architecture/governance_security.md",
            "docs/architecture/future_system_convergence.md",
            "docs/architecture/parser_ui_cli.md",
            "docs/plans/operational-core/source_synthesis.md",
            "docs/plans/operational-core/parser-mvp/plan.md",
            "docs/plans/parser-extraction/parser-mvp/plan.md",
            "docs/plans/parser-extraction/parser-mvp/adjacent-parser-and-inspection-location-review.md",
            "docs/contracts/parser_result_v1.md",
            "docs/contracts/eva_export_contract.md",
            "docs/contracts/eva_export_contract_v1.md",
            "docs/contracts/storage_adapter_contract.md",
            "docs/contracts/storage_adapter_contract_v1.md",
            "docs/contracts/work_item_contract_v1.md",
            "docs/contracts/provider_principal_config_contract_v1.md",
            "docs/contracts/review_audit_event_contract_v1.md",
            "docs/contracts/evidence_package_contract_v1.md",
            "docs/contracts/extraction_adapter_contract_v1.md",
            "docs/decisions/0003-operational-core-mvp.md",
            "docs/decisions/0004-ground-up-compatible-parser-rebuild.md",
            "docs/decisions/0005-shared-internal-app-local-accounts.md",
            "docs/decisions/0006-box-package-before-live-upload.md",
            "docs/decisions/0007-deterministic-first-parser.md",
            "docs/decisions/0008-private-real-corpus-only.md",
            "docs/decisions/options/ui_platform_options.md",
            "docs/decisions/options/backend_stack_options.md",
            "docs/decisions/options/state_store_options.md",
            "docs/decisions/options/cloud_document_intelligence_options.md",
            "docs/plans/operational-core/tickets/backlog_index.md",
            "docs/plans/operational-core/tickets/p0-foundation.md",
            "docs/plans/operational-core/tickets/p1-operational-core-mvp.md",
            "docs/plans/operational-core/tickets/p2-parser-hardening-provider-parity.md",
            "docs/plans/operational-core/tickets/p3-integrations-storage-eva-intake.md",
            "docs/plans/operational-core/tickets/p4-intelligence-engineer-communications.md",
            "docs/plans/operational-core/tickets/p5-platform-expansion.md",
            "docs/reference/data/provider_coverage_matrix.md",
            "docs/reference/originalplanning_index.md",
            "docs/reference/_index.md",
            "docs/security/source_safety_review.md",
            "docs/security/data_map.md",
            "docs/security/dpia_vendor_governance.md",
            "docs/security/vendor_register.md",
            "docs/security/api_security_standard.md",
            "docs/operations/release_and_rollback.md",
            "docs/operations/monitoring_runbooks.md",
            "docs/operations/runbooks/outlook-intake-stopped.md",
            "docs/operations/runbooks/box-upload-failure.md",
            "docs/operations/runbooks/eva-rejected-payload.md",
            "src/ccc_parser/core.py",
            "src/ccc_parser/cli.py",
            "src/ccc_parser/ui/app.py",
            "tests/test_scaffold_contracts.py",
            "docs/plans/operational-core/archived_plans/superseded/.gitkeep",
        ]

        GENERATED_PACKS = [
            "ce_phase4_agents_reviewed_plan",
            "ce_system_plans_enhanced",
            "cedocumentmapper_rebuild_plan_pack_all_zips",
            "collision_engineers_ai_tools_plans_markdown",
            "collision_engineers_bulk_data_research_pack",
            "originalplans_output",
            "phase7_expanded_markdown_plan",
            "testprojectcontext",
        ]

        RESEARCH_DOCS = [
            "docs/research/gptdeepresearch.md",
            "docs/research/gptevadeepresearch.md",
            "docs/research/siderpdf.md",
        ]

        PLANNED_WORKSPACE_TERMS = [
            "docs/plans/unified-platform/",
            "docs/plans/automation-centre/",
            "docs/plans/parser-extraction/",
            "docs/plans/case-workflow-state/",
            "docs/plans/provider-principal-config/",
            "docs/plans/intake-storage-integrations/",
            "docs/plans/evidence-estimate-review/",
            "docs/plans/vehicle-valuation-data/",
            "docs/plans/engineer-communications/",
            "docs/plans/ai-agents/",
            "docs/plans/mcp-and-tooling/",
            "docs/plans/agent-skills/",
            "docs/plans/ai-platform-tools/",
            "docs/plans/user-experience-interfaces/",
            "docs/plans/finance-billing/",
            "docs/plans/governance-security/",
            "docs/plans/operations-quality/",
            "docs/plans/analytics-data-platform/",
            "docs/plans/external-platform-partners/",
            "docs/plans/product-business/",
        ]

        PROVIDER_PRESETS = [
            "ALISON",
            "ALS",
            "AMS",
            "AX",
            "BC",
            "BLACK",
            "CNX (Engineers)",
            "DFD",
            "EVA (Engineers)",
            "FW (Garage)",
            "FW (Solicitor)",
            "HDUK",
            "KBS",
            "KERR",
            "KMR",
            "MP (Branded)",
            "MP (Simple)",
            "OAK",
            "PCH (Lawshield)",
            "PCH (Performance)",
            "QCL",
            "QDOS",
            "RJS",
            "SBL",
            "SWAN",
            "TEN",
        ]

        PARSER_PLAN_TERMS = [
            "PDF",
            "DOCX",
            "DOC",
            "MSG",
            "EML",
            "images",
            "batch",
            "PyMuPDF",
            "pdfplumber",
            "pypdf",
            "OCR only",
            "DD/MM/YYYY",
            "six-line",
            "EVA JSON field order",
            "UI/CLI",
            "deterministic-first",
            "private real corpus",
        ]

        PLAN_METADATA_TERMS = [
            "Status:",
            "Owner:",
            "Created:",
            "Last reviewed:",
            "Source links:",
            "Roadmap milestone:",
            "Dependencies:",
            "Expected outputs:",
            "Acceptance criteria:",
            "Verification required:",
            "Archive target:",
            "Supersedes:",
            "Superseded-by:",
        ]

        FALLBACK_IGNORED_MANIFEST_PARTS = {
            ".git",
            ".obsidian",
            ".pytest_cache",
            "__pycache__",
            ".mypy_cache",
            ".ruff_cache",
            ".venv",
            "venv",
            "env",
            "outputs",
            "tmp",
            "dist",
            "build",
            "node_modules",
        }
        FALLBACK_IGNORED_MANIFEST_NAMES = {".DS_Store", "Thumbs.db"}
        FALLBACK_IGNORED_MANIFEST_SUFFIXES = {".pyc", ".pyo", ".pyd", ".log", ".tmp", ".bak"}


        def require(condition: bool, message: str) -> None:
            if not condition:
                raise SystemExit(message)


        def read_doc(path: str) -> str:
            return (ROOT / path).read_text(encoding="utf-8")


        def require_terms(text: str, terms: list[str], context: str) -> None:
            missing = [term for term in terms if term not in text]
            require(not missing, f"{context} missing terms: {missing}")


        def is_fallback_ignored_manifest_artifact(path: Path) -> bool:
            parts = path.relative_to(ROOT).parts
            if set(parts) & FALLBACK_IGNORED_MANIFEST_PARTS:
                return True
            if any(part.endswith(".egg-info") for part in parts):
                return True
            name = path.name
            if name == ".env.example":
                return False
            if name == ".env" or name.startswith(".env."):
                return True
            if name in FALLBACK_IGNORED_MANIFEST_NAMES:
                return True
            return any(name.endswith(suffix) for suffix in FALLBACK_IGNORED_MANIFEST_SUFFIXES)


        def manifest_candidate_paths() -> list[str]:
            try:
                result = subprocess.run(
                    ["git", "ls-files", "--cached", "--others", "--exclude-standard", "-z"],
                    cwd=ROOT,
                    check=True,
                    capture_output=True,
                    text=True,
                )
            except (OSError, subprocess.CalledProcessError):
                candidates = []
                for path in ROOT.rglob("*"):
                    if path.is_file() and not is_fallback_ignored_manifest_artifact(path):
                        candidates.append(path.relative_to(ROOT).as_posix())
                return sorted(candidates)

            candidates = []
            for relative in result.stdout.split("\\0"):
                if not relative:
                    continue
                path = ROOT / relative
                if path.is_file():
                    candidates.append(Path(relative).as_posix())
            return sorted(set(candidates))


        def require_scope_guardrails() -> None:
            allowed_markers = [
                "out of scope",
                "must not",
                "does not plan",
                "do not plan",
                "do not add",
                "no personal injury",
                "no pi",
                "not plan",
                "confirm no",
            ]
            checked_roots = [
                "README.md",
                "AGENTS.md",
                "docs/architecture",
                "docs/contracts",
                "docs/decisions",
                "docs/plans/operational-core",
                "docs/plans",
                "docs/roadmap.md",
                "docs/security",
                "docs/plans/operational-core/tickets",
            ]
            files: list[Path] = []
            for item in checked_roots:
                path = ROOT / item
                if path.is_file():
                    files.append(path)
                elif path.is_dir():
                    files.extend(path.rglob("*.md"))
            bad_lines: list[str] = []
            for path in files:
                for line_number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
                    lowered = line.lower()
                    if "personal injury" in lowered or "kadoe" in lowered:
                        if not any(marker in lowered for marker in allowed_markers):
                            bad_lines.append(f"{path.relative_to(ROOT).as_posix()}:{line_number}: {line}")
            require(not bad_lines, "Personal injury/KADOE references must be explicit out-of-scope guardrails: " + repr(bad_lines[:5]))


        def main() -> int:
            missing = [path for path in REQUIRED_PATHS if not (ROOT / path).exists()]
            require(not missing, f"Missing required paths: {missing}")

            ambiguous_roots = [
                "originalplans_output",
                "ce_system_plans_enhanced",
                "ce_phase4_agents_reviewed_plan",
                "phase7_expanded_markdown_plan",
                "collision_engineers_ai_tools_plans_markdown",
                "collision_engineers_bulk_data_research_pack",
                "cedocumentmapper_rebuild_plan_pack_all_zips",
                "testprojectcontext",
            ]
            still_root = [path for path in ambiguous_roots if (ROOT / path).exists()]
            require(not still_root, f"Generated/reference packs still at repository root: {still_root}")

            legacy_roots = [
                "archive",
                "collisionrelateddocs",
                "docs/planning",
                "docs/tickets",
                "docs/normalized",
                "docs/data",
                "docs/reference/generated-packs",
            ]
            still_legacy = [path for path in legacy_roots if (ROOT / path).exists()]
            require(not still_legacy, f"Legacy documentation roots still exist: {still_legacy}")

            doubled_path_prefixes = [
                "docs/reference/raw/" + "docs/reference/raw",
                "docs/reference/data/jam_exports/" + "docs/reference/raw",
            ]
            doubled_path_hits: list[str] = []
            for path in ROOT.rglob("*"):
                if not path.is_file():
                    continue
                if set(path.relative_to(ROOT).parts) & {".git", ".obsidian", ".pytest_cache", "__pycache__"}:
                    continue
                rel_path = path.relative_to(ROOT).as_posix()
                if rel_path.startswith("docs/reference/raw/collisionrelateddocs/"):
                    continue
                if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".pdf", ".doc", ".docx", ".msg", ".xlsm", ".xlsx", ".jam", ".zip"}:
                    continue
                text = path.read_text(encoding="utf-8")
                if any(prefix in text for prefix in doubled_path_prefixes):
                    doubled_path_hits.append(rel_path)
            require(not doubled_path_hits, f"Doubled migrated path prefixes found: {doubled_path_hits[:10]}")

            providers = load_provider_presets(ROOT / "docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json")
            require(len(providers) == 26, f"Expected 26 providers, got {len(providers)}")
            provider_names = {provider.name for provider in providers}
            require(provider_names == set(PROVIDER_PRESETS), f"Provider preset names changed: {sorted(provider_names)}")

            manifest_rows = list(csv.DictReader((ROOT / "docs/source_manifest.csv").open(encoding="utf-8")))
            require(any(row["path"].startswith("docs/reference/raw/collisionrelateddocs/Instructions/") for row in manifest_rows), "Instruction corpus missing from manifest")
            require(any(row["normalized_companion"] for row in manifest_rows if row["path"].startswith("docs/reference/raw/collisionrelateddocs/")), "No normalized companions recorded")

            matrix_rows = list(csv.DictReader((ROOT / "docs/reference/data/provider_coverage_matrix.csv").open(encoding="utf-8")))
            require(any(row["code"] == "ACSP" and row["parser_covered"] == "no" for row in matrix_rows), "ACSP uncovered status missing")
            require(any(row["code"] == "WOODLANDS" and row["parser_covered"] == "no" for row in matrix_rows), "WOODLANDS uncovered status missing")

            plan_text = (ROOT / "docs/plans/initial-repo-setup/archived_plans/implemented/2026-05-23-implemented-initrepoplan.md").read_text(encoding="utf-8")
            require("Status: implemented" in plan_text, "Archived initrepoplan missing implemented status")

            synthesis_text = read_doc("docs/plans/operational-core/source_synthesis.md")
            require_terms(synthesis_text, GENERATED_PACKS, "Source synthesis")
            require_terms(synthesis_text, RESEARCH_DOCS, "Source synthesis")

            expansion_plan = read_doc("docs/plans/initial-repo-setup/documentation-scaffold/plans-folder-expansion-plan.md")
            require_terms(expansion_plan, PLANNED_WORKSPACE_TERMS, "Plans folder expansion plan")
            require_terms(
                expansion_plan,
                ["Deterministic automation", "AI agents", "MCP plans", "Agent skills"],
                "Plans folder expansion coverage",
            )
            require_terms(read_doc("docs/plans/_index.md"), PLANNED_WORKSPACE_TERMS, "Plans index planned workspaces")

            parser_plan = read_doc("docs/plans/parser-extraction/parser-mvp/plan.md")
            require_terms(parser_plan, PLAN_METADATA_TERMS, "Parser MVP plan metadata")
            require_terms(parser_plan, PROVIDER_PRESETS, "Parser MVP plan")
            require_terms(parser_plan, PARSER_PLAN_TERMS, "Parser MVP plan")

            for plan_file in [
                "docs/plans/operational-core/source_synthesis.md",
                "docs/plans/operational-core/tickets/backlog_index.md",
            ]:
                require_terms(read_doc(plan_file), PLAN_METADATA_TERMS, f"{plan_file} metadata")

            for ticket_file in [
                "docs/plans/operational-core/tickets/p0-foundation.md",
                "docs/plans/operational-core/tickets/p1-operational-core-mvp.md",
                "docs/plans/operational-core/tickets/p2-parser-hardening-provider-parity.md",
                "docs/plans/operational-core/tickets/p3-integrations-storage-eva-intake.md",
                "docs/plans/operational-core/tickets/p4-intelligence-engineer-communications.md",
                "docs/plans/operational-core/tickets/p5-platform-expansion.md",
            ]:
                ticket_text = read_doc(ticket_file)
                require_terms(
                    ticket_text,
                    PLAN_METADATA_TERMS,
                    ticket_file,
                )

            require_scope_guardrails()

            core = ParserCore(ROOT / "docs/reference/raw/collisionrelateddocs/Settings Backup/providers.json")
            require(len(core.providers) == 26, "ParserCore did not load providers")

            manifest_json_rows = json.loads((ROOT / "docs/source_manifest.json").read_text(encoding="utf-8"))
            require(len(manifest_json_rows) == len(manifest_rows), "Source manifest CSV/JSON row counts differ")
            manifest_paths = {row["path"] for row in manifest_rows}
            json_paths = {row["path"] for row in manifest_json_rows}
            require(json_paths == manifest_paths, "Source manifest CSV/JSON path sets differ")
            require(all(path in manifest_paths for path in REQUIRED_PATHS), "Source manifest missing required active docs")
            stale_manifest_paths = [path for path in manifest_paths if not (ROOT / path).exists()]
            require(not stale_manifest_paths, f"Source manifest includes missing files: {stale_manifest_paths[:10]}")
            missing_from_manifest: list[str] = []
            for relative in manifest_candidate_paths():
                if relative not in manifest_paths:
                    missing_from_manifest.append(relative)
            require(not missing_from_manifest, f"Files missing from source manifest: {missing_from_manifest[:10]}")
            print("Scaffold verification passed")
            return 0


        if __name__ == "__main__":
            raise SystemExit(main())
        """,
    )


def archive_plan() -> None:
    src = ROOT / "initrepoplan.md"
    if not src.exists():
        return
    content = src.read_text(encoding="utf-8")
    status = textwrap.dedent(
        f"""
        # Implemented State

        - Status: implemented
        - Implemented date: {TODAY}
        - Delivered summary: Repository hygiene docs, canonical documentation scaffold, source manifest, normalized companions/blockers, provider coverage matrix, parser UI/CLI scaffold, contracts, roadmap, and verification script were created.
        - Verification performed: `tools/verify_scaffold.py` plus targeted file/content checks.
        - Follow-up work moved elsewhere: Provider golden-rule implementation, production UI buildout, real EVA payload validation, live Box upload, cloud storage selection, and cloud OCR/document-intelligence vendor review.

        ---
        """
    ).strip()
    archive = ROOT / f"docs/plans/initial-repo-setup/archived_plans/implemented/{TODAY}-implemented-initrepoplan.md"
    write_text(archive, status + "\n\n" + content)
    src.unlink()


def main() -> int:
    organize_reference_packs()
    write_static_docs()
    write_generated_pack_index()
    write_parser_scaffold()

    initial_records = build_manifest()
    for record in initial_records:
        path = ROOT / record["path"]
        write_companion(path, record)

    write_provider_matrix()
    write_spreadsheet_flow_doc()
    archive_plan()

    final_records = build_manifest()
    write_manifest(final_records)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
